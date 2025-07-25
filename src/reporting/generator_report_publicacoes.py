# Importações necessárias
from langchain.prompts import PromptTemplate
from langchain.document_loaders import UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser
from langchain_groq import ChatGroq
from pydantic import BaseModel, Field
import pandas as pd
from typing import List, Optional
from langchain_core.exceptions import OutputParserException
import json

def preencher_publicacoes_(llm, pilares, objetivos, publico, posicionamento):

    # 1. Definição da estrutura de saída com Pydantic
    class Reel(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema principal do seu vídeo.")
        legenda: str = Field(description="A descrição da publicação.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        gancho: str = Field(description="A primeira frase ou cena para prender a atenção.")
        script_imagem: str = Field(description="Sequencia de imagens do vídeo.")
        script_audio: str = Field(description="Sequencia de audios do vídeo.")
        cta: str = Field(description="Call-to-action para engajamento")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")

    class Carrossel(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema do post.")
        legenda: str = Field(description="A descrição da publicação.")
        legenda_cta: str = Field(description="Chamada principal na legenda.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        capa: str = Field(description="Título extremamente chamativo que vai na primeira imagem.")
        lamina_cta: str = Field(description="A chamada para ação final no carrossel.")
        lamina_sequencia: str = Field(description="A sequencia de carrosseis.")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")
    
    class Imagem(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema do post.")
        legenda: str = Field(description="A descrição da publicação.")
        legenda_cta: str = Field(description="Chamada principal na legenda.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        imagem: str = Field(description="Descreva a imagem da publicação.")
        audio: Optional[str] = Field(description="O audio viral utilizado na publicação, caso aja")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")

    class Stories(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O assunto geral dos Stories do dia. Ex: Bastidores do Lançamento do Produto Y.")
        objetivo: str = Field(description="O que você quer alcançar? Ex: aumentar engajamento, Levar tráfego para o site, Vender.")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        capa: str = Field(description="Como você vai começar para prender a atenção? Descreva a imagem/vídeo e o texto do Primeiro Stories.")
        story_sequence: str = Field(description="Sequencia de Storys da arte.")
        elemento_interativo: str = Field(description="Qual sticker será usado na sequência? Use esta coluna como um checklist: Enquete, Caixa de Perguntas, Quiz, Link, Contagem Regressiva.")
        cta: str = Field(description="Qual a ação final esperada ao fim da sequência? Ex: Arrastar pra cima para se inscrever.")
        
    class PlanoDeConteudo(BaseModel):
        reels: List[Reel]
        carrossel: List[Carrossel]
        imagens: List[Imagem]
        stories: List[Stories]

    # Convert dicts to JSON strings as before
    objetivos_str = json.dumps(objetivos, indent=2, ensure_ascii=False)
    publico_str = json.dumps(publico, indent=2, ensure_ascii=False)
    posicionamento_str = json.dumps(posicionamento, indent=2, ensure_ascii=False)

    parser = PydanticOutputParser(pydantic_object=PlanoDeConteudo)
    todos_os_reels = []
    todos_os_carrosseis = []
    todas_as_imagens = []
    todos_os_stories = []
    historico_de_conteudo = {pilar['nome']: [] for pilar in pilares}

    for _ in range(2):
        
        for pilar in pilares:
            
            nome_do_pilar_atual = pilar['nome']
            print(f"Gerando conteúdo para o pilar: {nome_do_pilar_atual}...")
            conteudo_existente = historico_de_conteudo.get(nome_do_pilar_atual, [])
            historico_formatado = "\n".join(f"- {item}" for item in conteudo_existente)

            # ⬇️⬇️ START OF CHANGES ⬇️⬇️

            # 1. Use a regular multiline string, not an f-string.
            #    Use placeholders {objetivos}, {publico}, etc.
            prompt_template = """
            Você é um estrategista de conteúdo sênior, especialista em criar posts com conteúdo de valor.

            # INFORMAÇÕES ESTRATÉGICAS DO CLIENTE:
            ---
            ## OBJETIVOS:
            {objetivos}

            ## PÚBLICO-ALVO:
            {publico}

            ## POSICIONAMENTO E TOM DE VOZ:
            {posicionamento}
            ---

            # SUA TAREFA:
            Crie um plano de conteúdo para o Instagram focado EXCLUSIVAMENTE no pilar de conteúdo: "{pilar_atual}".
            O conteúdo deve ser EXTREMAMENTE CRIATIVO e DIVERSIFICADO.

            # REGRAS PARA CRIATIVIDADE (ESSENCIAL):
            1.  **Variedade de Ângulos**: Não crie posts com a mesma ideia. Para cada post, use um ângulo diferente (ex: tutorial passo a passo, desmistificar um mito, contar uma história, fazer uma pergunta polêmica, mostrar bastidores).
            2.  **Conexão com o Público**: O conteúdo deve resolver uma "dor" ou apelar para um "interesse" do público-alvo definido.
            3.  **Variação de CTA**: Não use sempre a mesma Call-to-Action. Varie entre "Comente sua opinião", "Salve este post", "Compartilhe com um amigo", "Faça uma pergunta nos comentários", etc.
            4.  **Alinhamento com o Tom de Voz**: O texto (legenda, título) deve seguir o tom de voz definido no posicionamento.
            5.  **Evitar Repetição**: O conteúdo abaixo já foi criado. NÃO REPITA estes temas ou ideias.
            ---
            ## CONTEÚDO JÁ CRIADO PARA ESTE PILAR:
            {historico}
            ---
            
            # INSTRUÇÕES DE GERAÇÃO:
            - Gere 6 posts de Reels, 3 de Carrossel, 3 de Imagem e 9 sequências de Stories.
            - Sua resposta DEVE SER APENAS o objeto JSON, sem nenhum texto antes ou depois.

            {format_instructions}
            """

            prompt = PromptTemplate(
                template=prompt_template,
                # 2. Define all the variables the template will need.
                input_variables=["objetivos", "publico", "posicionamento", "pilar_atual", "historico"],
                partial_variables={"format_instructions": parser.get_format_instructions()}
            )
            
            chain = prompt | llm | parser
            chain_with_retries = chain.with_retry(
                retry_if_exception_type=(OutputParserException, ValueError),
                stop_after_attempt=3
            )

            try:
                # 3. Pass all the variables in the .invoke() call.
                resultado_pilar = chain_with_retries.invoke({
                    "objetivos": objetivos_str,
                    "publico": publico_str,
                    "posicionamento": posicionamento_str,
                    "pilar_atual": nome_do_pilar_atual,
                    "historico": historico_formatado
                })
                # Adiciona os resultados às listas principais
                todos_os_reels.extend(resultado_pilar.reels)
                todos_os_carrosseis.extend(resultado_pilar.carrossel)
                todas_as_imagens.extend(resultado_pilar.imagens)
                todos_os_stories.extend(resultado_pilar.stories)
                print(f"Sucesso ao gerar conteúdo para {nome_do_pilar_atual}!")

                for item in resultado_pilar.reels:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.carrossel:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.imagens:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.stories:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)            

            except Exception as e:
                print(f"Falha ao processar o pilar {nome_do_pilar_atual}. Erro: {e}")

    # 6. Conversão para DataFrame e exportação (usando as listas acumuladas)
    df_dados_reels = pd.DataFrame([p.dict() for p in todos_os_reels])
    df_dados_carrossel = pd.DataFrame([p.dict() for p in todos_os_carrosseis])
    df_dados_imagem = pd.DataFrame([p.dict() for p in todas_as_imagens])
    df_dados_stories = pd.DataFrame([p.dict() for p in todos_os_stories])

    # Nome do arquivo Excel que terá todas as abas
    arquivo_saida = r"reports\publicações.xlsx"

    # Use o pd.ExcelWriter para criar e gerenciar o arquivo
    with pd.ExcelWriter(arquivo_saida, engine='openpyxl') as writer:
        # Escreve cada DataFrame em uma aba específica
        df_dados_reels.to_excel(writer, sheet_name='Reels', index=False)
        df_dados_carrossel.to_excel(writer, sheet_name='Carrossel', index=False)
        df_dados_imagem.to_excel(writer, sheet_name='Imagem', index=False)
        df_dados_stories.to_excel(writer, sheet_name='Stories', index=False)

    print(f"Relatório 'Publicações' salvo com sucesso em: {arquivo_saida}")

from langchain_community.document_loaders import UnstructuredWordDocumentLoader # CORREÇÃO: Nova importação
from config import settings
from langchain.llms import OpenAI
from langchain_community.document_loaders import TextLoader # Adicionado, se necessário
from langchain.text_splitter import CharacterTextSplitter # Adicionado, se necessário
from docx import Document
from docx.shared import Inches
import os

# Certifique-se de que o LLM é passado para a função

def preencher_publicacoes_(llm, pilares, objetivos, publico, posicionamento):
    try:
        # Carrega o documento modelo
        # CORREÇÃO: Usando o caminho de settings.py
        doc = Document(settings.TEMPLATE_PATH) # Ajustar para o seu caminho real do template

        # Texto para o LLM
        prompt_text = f"""
        Com base nos seguintes pilares de conteúdo, objetivos, público-alvo e posicionamento, gere 5 sugestões detalhadas de posts para redes sociais. Cada sugestão deve incluir:
        - Tema:
        - Tipo de Conteúdo: (ex: Foto, Vídeo, Carrossel, Reels, Stories)
        - Legenda:
        - Chamada para Ação (CTA):
        - Hashtags:
        - Objetivo Associado:

        Pilares de Conteúdo: {pilares}
        Objetivos: {objetivos}
        Público-alvo: {publico}
        Posicionamento: {posicionamento}
        """

        # Gerar sugestões de posts usando o LLM
        # Ajuste a chamada ao LLM conforme a sua integração
        # Aqui, estamos usando o LLM diretamente para gerar o texto
        print("Gerando sugestões de publicações com o LLM...")
        suggestions_raw = llm.invoke(prompt_text).content # Supondo que .invoke() é o método para gerar texto
        print("Sugestões geradas pelo LLM.")

        # Exemplo de como você pode tentar parsear as sugestões em um formato estruturado
        # Para um parsing mais robusto, você usaria PydanticOutputParser e um schema
        suggestions_list = suggestions_raw.strip().split('\n\n') # Divide por blocos de sugestão

        # Encontre o marcador no documento
        for paragraph in doc.paragraphs:
            if '[SUGESTOES_DE_PUBLICACAO]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[SUGESTOES_DE_PUBLICACAO]', '')
                for suggestion in suggestions_list:
                    if suggestion.strip():
                        # Adiciona cada sugestão como um novo parágrafo
                        run = paragraph.add_run(suggestion + '\n\n')
                break # Sai do loop após encontrar e preencher

        # Salva o documento atualizado
        # CORREÇÃO: Usando o caminho de settings.py
        doc.save(settings.REPORTS_PATH / "sugestoes_publicacoes.docx")
        print(f"Relatório de publicações salvo em: {settings.REPORTS_PATH / 'sugestoes_publicacoes.docx'}")

    except FileNotFoundError:
        print(f"Erro: O arquivo de template não foi encontrado em {settings.TEMPLATE_PATH}")
    except Exception as e:
        print(f"Ocorreu um erro ao gerar o relatório de publicações: {e}")
        raise # Re-lança o erro

def preencher_publicacoes(llm, pilares, objetivos, publico, posicionamento):

    # 1. Definição da estrutura de saída com Pydantic
    class Reel(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema principal do seu vídeo.")
        legenda: str = Field(description="A descrição da publicação.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        gancho: str = Field(description="A primeira frase ou cena para prender a atenção.")
        script_imagem: str = Field(description="Sequencia de imagens do vídeo.")
        script_audio: str = Field(description="Sequencia de audios do vídeo.")
        cta: str = Field(description="Call-to-action para engajamento")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")

    class Carrossel(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema do post.")
        legenda: str = Field(description="A descrição da publicação.")
        legenda_cta: str = Field(description="Chamada principal na legenda.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        capa: str = Field(description="Título extremamente chamativo que vai na primeira imagem.")
        lamina_cta: str = Field(description="A chamada para ação final no carrossel.")
        lamina_sequencia: str = Field(description="A sequencia de carrosseis.")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")
    
    class Imagem(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O tema do post.")
        legenda: str = Field(description="A descrição da publicação.")
        legenda_cta: str = Field(description="Chamada principal na legenda.")
        objetivo: str = Field(description="O que você quer alcançar?")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        imagem: str = Field(description="Descreva a imagem da publicação.")
        audio: Optional[str] = Field(description="O audio viral utilizado na publicação, caso aja")
        hashtags: str = Field(description="Hashtags relevantes separadas por vírgula")

    class Stories(BaseModel):
        pilar: str = Field(description="O pilar da publicação")
        titulo: str = Field(description="O assunto geral dos Stories do dia. Ex: Bastidores do Lançamento do Produto Y.")
        objetivo: str = Field(description="O que você quer alcançar? Ex: aumentar engajamento, Levar tráfego para o site, Vender.")
        publico: str = Field(description="Para quem este vídeo é direcionado?")
        capa: str = Field(description="Como você vai começar para prender a atenção? Descreva a imagem/vídeo e o texto do Primeiro Stories.")
        story_sequence: str = Field(description="Sequencia de Storys da arte.")
        elemento_interativo: str = Field(description="Qual sticker será usado na sequência? Use esta coluna como um checklist: Enquete, Caixa de Perguntas, Quiz, Link, Contagem Regressiva.")
        cta: str = Field(description="Qual a ação final esperada ao fim da sequência? Ex: Arrastar pra cima para se inscrever.")
        
    class PlanoDeConteudo(BaseModel):
        reels: List[Reel]
        carrossel: List[Carrossel]
        imagens: List[Imagem]
        stories: List[Stories]

    # Convert dicts to JSON strings as before
    objetivos_str = json.dumps(objetivos, indent=2, ensure_ascii=False)
    publico_str = json.dumps(publico, indent=2, ensure_ascii=False)
    posicionamento_str = json.dumps(posicionamento, indent=2, ensure_ascii=False)

    parser = PydanticOutputParser(pydantic_object=PlanoDeConteudo)
    todos_os_reels = []
    todos_os_carrosseis = []
    todas_as_imagens = []
    todos_os_stories = []
    historico_de_conteudo = {pilar['nome']: [] for pilar in pilares}

    for _ in range(2):
        
        for pilar in pilares:
            
            nome_do_pilar_atual = pilar['nome']
            print(f"Gerando conteúdo para o pilar: {nome_do_pilar_atual}...")
            conteudo_existente = historico_de_conteudo.get(nome_do_pilar_atual, [])
            historico_formatado = "\n".join(f"- {item}" for item in conteudo_existente)

            # ⬇️⬇️ START OF CHANGES ⬇️⬇️

            # 1. Use a regular multiline string, not an f-string.
            #    Use placeholders {objetivos}, {publico}, etc.
            prompt_template = """
            Você é um estrategista de conteúdo sênior, especialista em criar posts virais e de alto engajamento para o Instagram.

            # INFORMAÇÕES ESTRATÉGICAS DO CLIENTE:
            ---
            ## OBJETIVOS:
            {objetivos}

            ## PÚBLICO-ALVO:
            {publico}

            ## POSICIONAMENTO E TOM DE VOZ:
            {posicionamento}
            ---

            # SUA TAREFA:
            Crie um plano de conteúdo para o Instagram focado EXCLUSIVAMENTE no pilar de conteúdo: "{pilar_atual}".
            O conteúdo deve ser EXTREMAMENTE CRIATIVO e DIVERSIFICADO.

            # REGRAS PARA CRIATIVIDADE (ESSENCIAL):
            1.  **Variedade de Ângulos**: Não crie posts com a mesma ideia. Para cada post, use um ângulo diferente (ex: tutorial passo a passo, desmistificar um mito, contar uma história, fazer uma pergunta polêmica, mostrar bastidores).
            2.  **Conexão com o Público**: O conteúdo deve resolver uma "dor" ou apelar para um "interesse" do público-alvo definido.
            3.  **Variação de CTA**: Não use sempre a mesma Call-to-Action. Varie entre "Comente sua opinião", "Salve este post", "Compartilhe com um amigo", "Faça uma pergunta nos comentários", etc.
            4.  **Alinhamento com o Tom de Voz**: O texto (legenda, título) deve seguir o tom de voz definido no posicionamento.
            5.  **Evitar Repetição**: O conteúdo abaixo já foi criado. NÃO REPITA estes temas ou ideias.
            ---
            ## CONTEÚDO JÁ CRIADO PARA ESTE PILAR:
            {historico}
            ---
            
            # INSTRUÇÕES DE GERAÇÃO:
            - Gere 6 posts de Reels, 3 de Carrossel, 3 de Imagem e 9 sequências de Stories.
            - Sua resposta DEVE SER APENAS o objeto JSON, sem nenhum texto antes ou depois.

            {format_instructions}
            """

            prompt = PromptTemplate(
                template=prompt_template,
                # 2. Define all the variables the template will need.
                input_variables=["objetivos", "publico", "posicionamento", "pilar_atual", "historico"],
                partial_variables={"format_instructions": parser.get_format_instructions()}
            )
            
            chain = prompt | llm | parser
            chain_with_retries = chain.with_retry(
                retry_if_exception_type=(OutputParserException, ValueError),
                stop_after_attempt=3
            )

            try:
                # 3. Pass all the variables in the .invoke() call.
                resultado_pilar = chain_with_retries.invoke({
                    "objetivos": objetivos_str,
                    "publico": publico_str,
                    "posicionamento": posicionamento_str,
                    "pilar_atual": nome_do_pilar_atual,
                    "historico": historico_formatado
                })
                # Adiciona os resultados às listas principais
                todos_os_reels.extend(resultado_pilar.reels)
                todos_os_carrosseis.extend(resultado_pilar.carrossel)
                todas_as_imagens.extend(resultado_pilar.imagens)
                todos_os_stories.extend(resultado_pilar.stories)
                print(f"Sucesso ao gerar conteúdo para {nome_do_pilar_atual}!")

                for item in resultado_pilar.reels:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.carrossel:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.imagens:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)
                for item in resultado_pilar.stories:
                    historico_de_conteudo[nome_do_pilar_atual].append(item.titulo)            

            except Exception as e:
                print(f"Falha ao processar o pilar {nome_do_pilar_atual}. Erro: {e}")

    # 6. Conversão para DataFrame e exportação (usando as listas acumuladas)
    df_dados_reels = pd.DataFrame([p.dict() for p in todos_os_reels])
    df_dados_carrossel = pd.DataFrame([p.dict() for p in todos_os_carrosseis])
    df_dados_imagem = pd.DataFrame([p.dict() for p in todas_as_imagens])
    df_dados_stories = pd.DataFrame([p.dict() for p in todos_os_stories])

    # Nome do arquivo Excel que terá todas as abas
    arquivo_saida = r"reports\publicações.xlsx"

    # Use o pd.ExcelWriter para criar e gerenciar o arquivo
    with pd.ExcelWriter(arquivo_saida, engine='openpyxl') as writer:
        # Escreve cada DataFrame em uma aba específica
        df_dados_reels.to_excel(writer, sheet_name='Reels', index=False)
        df_dados_carrossel.to_excel(writer, sheet_name='Carrossel', index=False)
        df_dados_imagem.to_excel(writer, sheet_name='Imagem', index=False)
        df_dados_stories.to_excel(writer, sheet_name='Stories', index=False)

    print(f"Relatório 'Publicações' salvo com sucesso em: {arquivo_saida}")
