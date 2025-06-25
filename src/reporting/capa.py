from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerarCapa(document, titulo, cliente, autor, data, resumo_profissional):
    """
    Gera um documento Word com capa e resumo profissional para uma análise de concorrentes.

    Args:
        titulo (str): O título do relatório.
        cliente (str): O nome do cliente para quem a análise foi preparada.
        autor (str): O nome do autor ou da agência/consultoria.
        data (str): A data de publicação do relatório.
        resumo_profissional (str): O texto do resumo profissional.
    """
    # --- Seção da Capa ---

    # Título do Relatório
    titulo_principal = document.add_heading(titulo, level=1)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_principal.runs:
        run.font.size = Pt(24)
        run.bold = True

    # Adiciona espaçamento após o título
    document.add_paragraph()
    document.add_paragraph()

    # Informações do Cliente e Autor
    p_info = document.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f'Preparado para:\n{cliente}\n\n').bold = True
    p_info.add_run(f'Análise por:\n{autor}\n\n')

    # Data
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.add_run(data)

    # Adiciona quebra de página
    document.add_page_break()

    # --- Seção do Resumo Profissional ---

    # Título do Resumo
    titulo_resumo = document.add_heading('Resumo Profissional', level=2)
    for run in titulo_resumo.runs:
        run.font.size = Pt(16)
        run.bold = True

    # Corpo do Resumo
    p_resumo = document.add_paragraph(resumo_profissional)
    p_resumo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Salva o documento
    nome_arquivo = f"Analise_Concorrentes_Instagram_{cliente.replace(' ', '_')}.docx"
    document.save(nome_arquivo)
    print(f"Relatório '{nome_arquivo}' gerado com sucesso!")


if __name__ == '__main__':
    # --- Preencha as informações da sua análise aqui ---
    titulo_analise = "Análise de Concorrentes no Instagram"
    nome_cliente = "Cliente X"
    nome_autor = "Sua Agência / Seu Nome"
    data_analise = "Junho de 2025"
    texto_resumo = (
        "Este relatório apresenta uma análise competitiva do desempenho no Instagram, "
        "preparada para o Cliente X. O estudo foi conduzido entre maio e junho de 2025 e "
        "avaliou os perfis dos principais concorrentes: Concorrente A, Concorrente B e "
        "Concorrente C. A análise focou em métricas de engajamento, crescimento de seguidores, "
        "estratégia de conteúdo e frequência de postagens. Os resultados indicam que o "
        "Concorrente B lidera em engajamento, utilizando fortemente o formato de Reels com "
        "conteúdo gerado pelo usuário. O Concorrente A, por sua vez, demonstra um crescimento "
        "acelerado de seguidores através de parcerias com influenciadores. Identificou-se uma "
        "oportunidade para o Cliente X explorar o formato de Stories interativos (enquetes, "
        "caixas de perguntas), pouco utilizado pela concorrência. Recomenda-se a adoção de uma "
        "estratégia de conteúdo em vídeo mais robusta, similar à do Concorrente B, e a "
        "implementação de um calendário editorial que inclua colaborações estratégicas para "
        "aumentar o alcance e o engajamento do perfil."
    )

    gerar_analise_concorrentes(
        titulo_analise,
        nome_cliente,
        nome_autor,
        data_analise,
        texto_resumo
    )