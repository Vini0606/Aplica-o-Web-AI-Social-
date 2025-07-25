from docx import Document
from datetime import datetime
import os
import json
from docx import Document
from docx.text.paragraph import Paragraph
from datetime import datetime
import os
from docx.oxml.shared import OxmlElement
from docx import Document
from docx.text.paragraph import Paragraph
from copy import deepcopy
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm, RGBColor
from datetime import date
from config import settings

# =================================================================
# 헬 FUNÇÃO AUXILIAR PARA ESTILOS
# =================================================================
def definir_estilos(document):
    """
    Centraliza todas as modificações de estilo do documento para garantir consistência.
    """
    # 1. ALTERAR A FONTE PADRÃO (ESTILO 'NORMAL')
    estilo_normal = document.styles['Normal']
    fonte_normal = estilo_normal.font
    fonte_normal.name = 'Abadi'
    fonte_normal.size = Pt(11)

    # 2. ALTERAR ESTILO DO TÍTULO PRINCIPAL (NÍVEL 1)
    # Usado pela função `gerarCapa` e outros títulos de nível 1
    estilo_h1 = document.styles['Heading 1']
    fonte_h1 = estilo_h1.font
    fonte_h1.name = 'Abadi'
    fonte_h1.size = Pt(22)
    fonte_h1.bold = True
    fonte_h1.color.rgb = RGBColor(0x0A, 0x25, 0x40)
    paragrafo_h1 = estilo_h1.paragraph_format
    paragrafo_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_h1.space_before = Pt(12)
    paragrafo_h1.space_after = Pt(12)

    # 3. ALTERAR ESTILO DOS SUBTÍTULOS (NÍVEL 2)
    estilo_h2 = document.styles['Heading 2']
    fonte_h2 = estilo_h2.font
    fonte_h2.name = 'Abadi'
    fonte_h2.size = Pt(14)
    fonte_h2.bold = True
    fonte_h2.color.rgb = RGBColor(0x0A, 0x25, 0x40)  # Azul corporativo
    paragrafo_h2 = estilo_h2.paragraph_format
    paragrafo_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo_h2.space_before = Pt(18)
    paragrafo_h2.space_after = Pt(6)

    # 4. ALTERAR ESTILO DOS TÍTULOS DE TÓPICOS (NÍVEL 3)
    estilo_h3 = document.styles['Heading 3']
    fonte_h3 = estilo_h3.font
    fonte_h3.name = 'Abadi'
    fonte_h3.size = Pt(12)
    fonte_h3.bold = True
    paragrafo_h3 = estilo_h3.paragraph_format
    paragrafo_h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo_h3.space_before = Pt(10)
    paragrafo_h3.space_after = Pt(4)

# =================================================================
# 헬 FUNÇÃO PARA GERAR A CAPA
# =================================================================
def gerarCapa(document, titulo, cliente, autor, data):
    """
    Gera a capa do documento Word. A formatação do título vem do estilo 'Heading 1'.
    """
    # 1. Adicionar a logo da empresa na capa
    p_logo = document.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    # CORREÇÃO: Usando um valor sensato para a largura, como 6 cm.
    run_logo.add_picture(str(settings.LOGO_PATH), width=Cm(6))

    # 2. Título do Relatório (usa o estilo 'Heading 1' definido anteriormente)
    document.add_heading(titulo, level=1)

    # 3. Informações do Cliente e Autor
    document.add_paragraph() # Espaçamento
    p_info = document.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f'Preparado para:\n').bold = True
    p_info.add_run(f'{cliente}\n\n')
    p_info.add_run(f'Análise por:\n').bold = True
    p_info.add_run(f'{autor}\n\n')

    # 4. Data
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.add_run(data)
    document.add_page_break()

# =================================================================
# 헬 FUNÇÃO PRINCIPAL PARA PREENCHER O PLANO
# =================================================================
def preencher_plano_marketing(
    brief_data,
    caminho_saida,
    nome_empresa,
    objetivos,
    persona,
    pilares_conteudo,
    posicionamento,
    calendario=[]
):
    """
    Gera o documento completo do plano de marketing de conteúdo.
    """
    doc = Document()

    # 1. APLICA TODOS OS ESTILOS AO DOCUMENTO ANTES DE COMEÇAR
    definir_estilos(doc)

    # 2. CONFIGURA O CABEÇALHO PARA TODAS AS PÁGINAS (EXCETO A CAPA)
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    p_header = header.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = p_header.add_run()
    run_header.add_picture(str(settings.LOGO_PATH), height=Cm(3.5))

    # 3. GERA A CAPA
    nome_cliente = brief_data['objetivos']['client_name']
    gerarCapa(
        document=doc,
        titulo="Plano de Marketing de Conteúdo para Instagram",
        cliente=nome_cliente,
        autor="Equipe Social Planner",
        data=f"{date.today().strftime('%A, %d de %B de %Y')}"
    )

    # 4. GERA O CONTEÚDO DO RELATÓRIO USANDO OS ESTILOS DEFINIDOS
    doc.add_heading("📌 Objetivos de Marketing", level=2)
    for objetivo in objetivos:
        doc.add_paragraph(f"{objetivo}", style='List Bullet')

    doc.add_heading("🎯 Público-alvo", level=2)
    for chave, valor in persona.items():
        p = doc.add_paragraph()
        p.add_run(f"{chave.replace('_', ' ').capitalize()}: ").bold = True
        p.add_run(str(valor) if not isinstance(valor, list) else ', '.join(valor))

    doc.add_heading("🎙️ Posicionamento e Tom de Voz", level=2)
    doc.add_paragraph(posicionamento['resumo_posicionamento'])

    doc.add_heading("📚 Pilares de Conteúdo", level=2)
    for pilar in pilares_conteudo:
        doc.add_heading(f"{pilar['nome']}", level=3)
        doc.add_paragraph(pilar['objetivo'])
        doc.add_paragraph("Exemplos de conteúdo:")
        for exemplo in pilar.get('exemplos', []):
            doc.add_paragraph(f"{exemplo}", style='List Bullet')

    doc.add_heading("🗓️ Calendário Editorial Sugerido", level=2)
    if calendario:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dia'
        hdr_cells[1].text = 'Pilar de Conteúdo'
        hdr_cells[2].text = 'Horário Sugerido'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        for item in calendario:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('dia', 'N/A')
            row_cells[1].text = item.get('pilar', 'N/A')
            row_cells[2].text = item.get('periodo', 'N/A')
    else:
        doc.add_paragraph("Nenhuma sugestão de calendário foi gerada.")

    # ... (Adicione outras seções como 'Formatos', 'Estratégia de Engajamento', 'KPIs' aqui)
    doc.add_heading("📈 Estratégia de Engajamento e Crescimento", level=2)
    doc.add_paragraph("Interação Proativa: Dedicar tempo para responder comentários e DMs.", style='List Bullet')
    doc.add_paragraph("Uso Estratégico de Hashtags: Misturar hashtags de nicho e de volume.", style='List Bullet')
    doc.add_paragraph("Call to Actions (CTAs) Claras: Incentivar ações como salvar, comentar, etc.", style='List Bullet')

    # 5. SALVA O DOCUMENTO FINAL
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)
    print(f"✅ Relatório gerado com sucesso: {caminho_saida}")

# =================================================================
# 🎯 BLOCO DE EXECUÇÃO PRINCIPAL
# =================================================================
if __name__ == "__main__":
    with open(settings.BRIEFING_JSON_PATH, 'r', encoding='utf-8') as arquivo_json:
        brief_data = json.load(arquivo_json)

    objetivo_principal = [brief_data['objetivos']['objetivo_principal']]
    objetivos_secundarios = brief_data['objetivos']['objetivo_secundario']
    list_objetivos = objetivo_principal + objetivos_secundarios

    preencher_plano_marketing(
        brief_data=brief_data,
        caminho_saida=settings.ESTRATEGIA_PATH,
        nome_empresa=brief_data['objetivos']['client_name'],
        objetivos=list_objetivos,
        persona=brief_data['publico'],
        pilares_conteudo=brief_data['pilares'],
        posicionamento=brief_data['posicionamento'],
        calendario=brief_data.get('calendario', [])
    )