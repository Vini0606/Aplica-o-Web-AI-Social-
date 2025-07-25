import io
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import gridspec
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from langchain_ollama import ChatOllama
from langchain_groq import ChatGroq
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from wordcloud import WordCloud
from datetime import date
import locale
import numpy as np
import time
from config import settings

try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    # Em sistemas Windows:
    try:
        locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil')
    except locale.Error:
        print("Locale 'pt_BR.UTF-8' ou 'Portuguese_Brazil' não encontrado no sistema.")
        # Fallback para o locale padrão do sistema
        locale.setlocale(locale.LC_ALL, '')

def safe_invoke(llm, prompt):
    try:
        prompt.encode("utf-8")  # Testa se o prompt tem caracteres válidos
        return llm.invoke(prompt)
    except UnicodeEncodeError as e:
        print("⚠️ Prompt contém caracteres inválidos. Tentando limpar...")
        # Remove caracteres malformados
        prompt_clean = prompt.encode("utf-8", "ignore").decode("utf-8", "ignore")
        return llm.invoke(prompt_clean)

# --- Funções Auxiliares de Estilização ---
def set_cell_shading(cell, hex_color: str):
    """Define a cor de fundo de uma célula da tabela.""" 
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading_elm) 
    except Exception as e:
        print(f"Não foi possível aplicar a cor de fundo à célula: {e}") 

def set_header_style(table):
    """Aplica cor de fundo cinza e negrito ao cabeçalho da tabela.""" 
    header_cells = table.rows[0].cells
    for cell in header_cells:
        set_cell_shading(cell, 'D9D9D9') # Cinza claro
        
        # Para aplicar negrito
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run()
        run.text = cell.text
        run.bold = True
        p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER

# --- Gráficos ---
def Secao_2_1_Figura1(client_name, df_profiles_posts):

    def plotarBarraMax(client_name, x_col, y_col, ax1, x_col_name, y_col_name):
        
        # Calcular Estatísticas
        top_10 = df_profiles_posts.groupby([y_col])[x_col].max().sort_values(ascending=True).dropna().tail(10).reset_index()
        
        # --- 4. Configuração do Primeiro Gráfico (Superior) ---
        # Preparar cores e rótulos para o primeiro gráfico
        cores1 = ['#1f77b4'] * 10 # Cor padrão para as barras
        
        if client_name in list(top_10['username']):
            try:
                # Encontra a posição (índice) do elemento X nos dados ordenados
                indice_x = top_10.loc[top_10[y_col] == client_name].index[0]
                cores1[indice_x] = '#000080' # Cor de destaque para o elemento X
                print('Cor alterada com Sucesso!')
            except KeyError:
                indice_x = -1 # Trata o caso do elemento não ser encontrado

        # Plota o gráfico de barras
        barras1 = ax1.barh(top_10[y_col], top_10[x_col], color=cores1)
        ax1.bar_label(barras1, fmt='%d', padding=5)
        ax1.set_title(f'{y_col_name} X {x_col_name}')
        ax1.set_xlabel(x_col_name)
        ax1.set_ylabel(y_col_name)

        return top_10

    # --- 3. Criação da Figura e dos Gráficos (Subplots) ---
    # Cria a figura com 2 linhas e 1 coluna de gráficos
    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(16, 5))
      
    top_10_followers = plotarBarraMax(client_name, 'followersCount', 'username', ax1, 'Qtd de Seguidores', 'Usuários')
    top_10_follows = plotarBarraMax(client_name, 'followsCount', 'username', ax2, 'Qtd Seguindo', 'Usuários')
    top_10_posts_count = plotarBarraMax(client_name, 'postsCount', 'username', ax3, 'Qtd de Posts', 'Usuários')

    dataframes = {
                    "followers": top_10_followers,
                    "follows": top_10_follows,
                    "posts_count": top_10_posts_count
    }

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # Para exibir a figura diretamente (se estiver em um ambiente interativo como Jupyter)
    plt.close()

    buffer.seek(0)

    print("Figura 'graficos_de_barras_destacados.png' gerada com sucesso.")

    return buffer, dataframes

def Secao_2_1_Figura2(df_profiles_posts):

    def plotarFiguraNColsPCA(df_original, df_cluster, fig):
            
        def plotarDispersao(fig):
            
            cluster_columns = [column for column in df_cluster.columns if column != 'Clusters (AutoClusterHPO)']
            
            # Redução de dimensionalidade com PCA
            pca = PCA(n_components=2)
            df_pca = pca.fit_transform(df_original[cluster_columns])
            df_original_copy['pca_x'] = df_pca[:, 0]
            df_original_copy['pca_y'] = df_pca[:, 1]
                
            # Visualização dos clusters com PCA
            axes_gs_principal = fig.add_subplot(gs_principal[0])
            sns.scatterplot(data=df_original_copy, x='pca_x', y='pca_y', hue='Clusters (AutoClusterHPO)', palette='deep', s=100, ax=axes_gs_principal) 
            axes_gs_principal.set_title(f" X ".join(cluster_columns))
            axes_gs_principal.set_xlabel('Componente Principal 1')
            axes_gs_principal.set_ylabel('Componente Principal 2')
            axes_gs_principal.grid(True)

            return df_original_copy[['pca_x', 'pca_y', 'Clusters (AutoClusterHPO)']]
        
        def plotarBarras(fig):
            
            df_original['Clusters (AutoClusterHPO)'] = 'Cluster ' + df_original['Clusters (AutoClusterHPO)'].astype(str)
                
            media_1 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[0])].mean().reset_index().sort_values(by=str(df_cluster.columns[0]), ascending=True)
            media_2 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[1])].mean().reset_index().sort_values(by=str(df_cluster.columns[1]), ascending=True)
            media_3 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[2])].mean().reset_index().sort_values(by=str(df_cluster.columns[2]), ascending=True)
                
                
            # Plotar Gráfico de Barra 1
            axes_gs_0 = fig.add_subplot(gs[0])
            axes_gs_0.set_title(f'Clusters X {str(df_cluster.columns[0])}')
            barras1 = axes_gs_0.barh(media_1['Clusters (AutoClusterHPO)'], media_1[str(df_cluster.columns[0])])
            axes_gs_0.bar_label(barras1, fmt='%d', padding=-45)
            axes_gs_0.set_xticklabels([])
            axes_gs_0.set_xlabel(f'Média de {str(df_cluster.columns[0])}')

            # Plotar Gráfico de Barra 2
            axes_gs_1 = fig.add_subplot(gs[1])
            axes_gs_1.set_title(f'Clusters X {str(df_cluster.columns[1])}')
            barras2 = axes_gs_1.barh(media_2['Clusters (AutoClusterHPO)'], media_2[str(df_cluster.columns[1])])
            axes_gs_1.bar_label(barras2, fmt='%d', padding=5)
            axes_gs_1.set_xticklabels([])
            axes_gs_1.set_xlabel(f'Média de {str(df_cluster.columns[1])}')
                
            # Plotar Gráfico de Barra 3
            axes_gs1_0 = fig.add_subplot(gs1[0])
            axes_gs1_0.set_title(f'Clusters X {str(df_cluster.columns[2])}')
            barras3 = axes_gs1_0.barh(media_3['Clusters (AutoClusterHPO)'], media_3[str(df_cluster.columns[2])])
            axes_gs1_0.bar_label(barras3, fmt='%d', padding=-30)
            axes_gs1_0.set_xticklabels([])
            axes_gs1_0.set_xlabel(f'Média de {str(df_cluster.columns[2])}')

            return media_1, media_2, media_3
    
        # Copiar df original
        df_original_copy = df_original.copy()
            
        gs_principal = gridspec.GridSpec(1, 3, figure=fig, width_ratios=[2, 1, 1], wspace=0.2)
        gs = gs_principal[1].subgridspec(2, 1, hspace=0.4)
        gs1 = gs_principal[2].subgridspec(2, 1, hspace=0.4)
            
        df_original_copy = plotarDispersao(fig)
        media_1, media_2, media_3 = plotarBarras(fig)

        dataframes = {'df_original_copy' : df_original_copy, 
                    'df_media_1' : media_1,
                    'df_media_2' : media_2,
                    'df_media_3' : media_3}

        return dataframes
    
    # Separar Bases 
    df_original = df_profiles_posts
    df_cluster = df_profiles_posts[['followersCount', 'followsCount', 'postsCount']].copy()

    # Inicializar e aplicar o AutoCluster
    autocluster_tool = AutoClusterHPO(max_evals_per_algo=100) 
    df_original['Clusters (AutoClusterHPO)'], model, config, score, algo_name = autocluster_tool.fit_predict(df_cluster)

    fig = plt.figure(figsize=(16, 5))
    
    dataframes = plotarFiguraNColsPCA(df_original, df_cluster, fig)

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
        
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)
        
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_1_Figura3(posts_df):

    def plotarNuvemPalavras():

        lista_unica = []

        for sublista in posts_df['hashtags']:
            if isinstance(sublista, list):
                for item in sublista:
                    lista_unica.append(item)

        # Seu texto aqui
        texto = " ".join(lista_unica)

        # Criar o objeto WordCloud
        nuvem_palavras = WordCloud(width=800, height=400, background_color="white").generate(texto)
        
        return nuvem_palavras, pd.DataFrame(lista_unica).value_counts()

    nuvem_palavras, df = plotarNuvemPalavras()

    # Exibir a imagem gerada
    plt.figure(figsize=(16, 8))
    plt.imshow(nuvem_palavras, interpolation='bilinear')
    plt.axis("off") # Remove os eixos

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
        
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)
        
    plt.close()

    buffer.seek(0)

    return buffer, df 

def Secao_2_2_Figura4(client_name, posts_df, df_profiles_posts):
 
    def plotarBarraSum(client_name, df, x_col, y_col, ax1, fmt, x_col_name, y_col_name):
        
        # Calcular Estatísticas
        top_10 = df.groupby([y_col])[x_col].sum().sort_values(ascending=True).tail(10).reset_index()
        
        # --- 4. Configuração do Primeiro Gráfico (Superior) ---
        # Preparar cores e rótulos para o primeiro gráfico
        cores1 = ['#1f77b4'] * 10 # Cor padrão para as barras
        
        if client_name in list(top_10[y_col]):
            try:
                # Encontra a posição (índice) do elemento X nos dados ordenados
                indice_x = top_10.loc[top_10[y_col] == client_name].index[0]
                cores1[indice_x] = '#000080' # Cor de destaque para o elemento X
                print('Cor alterada com Sucesso!')
            except KeyError:
                indice_x = -1 # Trata o caso do elemento não ser encontrado

        # Plota o gráfico de barras
        barras1 = ax1.barh(top_10[y_col], top_10[x_col], color=cores1)
        ax1.bar_label(barras1, fmt=fmt, padding=5)
        ax1.set_title(f'{y_col_name} X {x_col_name}')
        ax1.set_xlabel(x_col_name)
        ax1.set_ylabel(y_col_name)

        return top_10

    # --- 3. Criação da Figura e dos Gráficos (Subplots) ---
    # Cria a figura com 2 linhas e 1 coluna de gráficos
    fig, axes = plt.subplots(2, 3, figsize=(16, 5))
       
    top_10_likes_profiles = plotarBarraSum(client_name, df_profiles_posts, 'likesSum', 'username', axes[0, 0], '%d', 'Qtd de Curtidas', 'Usuário')
    top_10_comments_profiles = plotarBarraSum(client_name, df_profiles_posts, 'commentsSum', 'username', axes[0, 1], '%d', 'Qtd de Comentários', 'Usuário')
    top_10_perc_engaj_profiles = plotarBarraSum(client_name, df_profiles_posts, '% ENGAJAMENTO', 'username', axes[0, 2], '%.2f', 'Taxa de Engajamento', 'Usuário')
    top_10_likes_posts = plotarBarraSum(client_name, posts_df, 'likesCount', 'shortCode', axes[1, 0], '%d', 'Qtd de Curtidas', 'ShortCode (Publicação)')
    top_10_comments_posts = plotarBarraSum(client_name, posts_df, 'commentsCount', 'shortCode', axes[1, 1], '%d', 'Qtd de Comentários', 'ShortCode (Publicação)')
    top_10_engaj_posts = plotarBarraSum(client_name, posts_df, 'TOTAL ENGAJAMENTO', 'shortCode', axes[1, 2], '%d', 'Qtd de Engajamentos', 'ShortCode (Publicação)')

    dataframes = {
                    "top_10_likes_profiles": top_10_likes_profiles,
                    "top_10_comments_profiles": top_10_comments_profiles,
                    "top_10_perc_engaj_profiles": top_10_perc_engaj_profiles,
                    "top_10_likes_posts":top_10_likes_posts,
                    "top_10_comments_posts":top_10_comments_posts,
                    "top_10_engaj_posts":top_10_engaj_posts

    }

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # Para exibir a figura diretamente (se estiver em um ambiente interativo como Jupyter)
    plt.close()

    buffer.seek(0)

    print("Figura 'graficos_de_barras_destacados.png' gerada com sucesso.")

    return buffer, dataframes

def Secao_2_2_Figura5(dados_pivot_count, dados_pivot_total):   

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax, x_column_name, y_column_name):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(y_column_name, fontsize=12)
        ax.set_ylabel(x_column_name, fontsize=12)
        ax.set_title(f'{y_column_name} X {x_column_name}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    dataframes = {'Type': dados_pivot_count,
                  'total': dados_pivot_total}
    
    print(dados_pivot_count)
    
    print(dados_pivot_total)

    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(2, 1, figsize=(15, 10))
    
    plotarBarrasAgrupadas(dados_pivot_count, 'countType', 'ownerUsername', 'Type', axes[0], 'Qtd de Posts', 'Usuários')
    plotarBarrasAgrupadas(dados_pivot_total, 'ENGAJAMENTO TOTAL', 'ownerUsername', 'Type', axes[1], 'Qtd de Engajamentos', 'Usuários')

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_2_Figura6(dados_pivot_likes, dados_pivot_comments):

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax, x_column_name, y_column_name):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(y_column_name, fontsize=12)
        ax.set_ylabel(x_column_name, fontsize=12)
        ax.set_title(f'{y_column_name} X {x_column_name}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    dataframes = {'likes': dados_pivot_likes, 
                  'comments': dados_pivot_comments}
    
    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(2, 1, figsize=(15, 10))
    
    plotarBarrasAgrupadas(dados_pivot_likes, 'likes', 'ownerUsername', 'Type', axes[0], 'Qtd de Curtidas', 'Usuários')
    plotarBarrasAgrupadas(dados_pivot_comments, 'comments', 'ownerUsername', 'Type', axes[1], 'Qtd de Comentários', 'Usuários')

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_3_Figura7(client_name, periodo_df, dias_df):

    def plotarBarraSum(top_10, x_col, y_col, ax1, fmt, x_col_name, y_col_name):
        
        # --- 4. Configuração do Primeiro Gráfico (Superior) ---
        # Preparar cores e rótulos para o primeiro gráfico
        cores1 = ['#1f77b4'] * 10 # Cor padrão para as barras

        # Plota o gráfico de barras
        barras1 = ax1.bar(top_10.index, top_10.values, color=cores1)
        ax1.bar_label(barras1, fmt=fmt, padding=5)
        ax1.set_title(f'{y_col_name} X {x_col_name}')
        ax1.set_xlabel(x_col_name)
        ax1.set_ylabel(y_col_name)

        return top_10

    # --- 3. Criação da Figura e dos Gráficos (Subplots) ---
    # Cria a figura com 2 linhas e 1 coluna de gráficos
    fig, axes = plt.subplots(nrows=2, ncols=1, figsize=(12, 10), gridspec_kw={'height_ratios': [1, 2]})
      
    periodos_df = plotarBarraSum(periodo_df, 'Periodo do Dia', 'Count', axes[0], '%d', 'Periodo do Dia', 'Qtd de Posts')
    dias_df = plotarBarraSum(dias_df, 'Dia da Semana', 'Count', axes[1], '%d', 'Dias da Semana', 'Qtd de Posts')

    dataframes = {
                    "periodos": periodos_df,
                    "dias": dias_df
    }

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # Para exibir a figura diretamente (se estiver em um ambiente interativo como Jupyter)
    plt.close()

    buffer.seek(0)

    print("Figura 'graficos_de_barras_destacados.png' gerada com sucesso.")

    return buffer, dataframes

def Secao_2_3_Figura8(dados_pivot_periodos, dados_pivot_dias):

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax, x_column_name, y_column_name):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(x_column_name, fontsize=12)
        ax.set_ylabel(y_column_name, fontsize=12)
        ax.set_title(f'{y_column_name} X {x_column_name}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        
    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(figsize=(15, 10))

    # Lista com a nova ordem desejada para as cidades
    ordem_periodos = ['Manhã', 'Tarde', 'Noite', 'Madrugada']
    ordem_dias = ['Segunda-feira', 'Terça-feira', 'Quarta-feira', 'Quinta-feira', 'Sexta-feira', 'Sábado', 'Domingo']

    periodo_df = dados_pivot_periodos.reindex(ordem_periodos)
    dias_df = dados_pivot_dias.reindex(ordem_dias)
    
    plotarBarrasAgrupadas(dias_df, 'Dias da Semana', 'Contagem', 'Type', axes, 'Dias de Semana', 'Qtd de Posts')

    dataframes = {'periodos': dados_pivot_periodos,
                  'dias': dados_pivot_dias}
    

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_3_Figura9(dados_pivot_periodos, dados_pivot_dias):

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax, x_column_name, y_column_name):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(x_column_name, fontsize=12)
        ax.set_ylabel(y_column_name, fontsize=12)
        ax.set_title(f'{y_column_name} X {x_column_name}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        
    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(figsize=(15, 10))

    # Lista com a nova ordem desejada para as cidades
    ordem_periodos = ['Manhã', 'Tarde', 'Noite', 'Madrugada']
    ordem_dias = ['Segunda-feira', 'Terça-feira', 'Quarta-feira', 'Quinta-feira', 'Sexta-feira', 'Sábado', 'Domingo']

    periodo_df = dados_pivot_periodos.reindex(ordem_periodos)
    dias_df = dados_pivot_dias.reindex(ordem_dias)
    
    plotarBarrasAgrupadas(periodo_df, 'Dias da Semana', 'Contagem', 'Type', axes, 'Periodos do Dia', 'Qtd de Posts')

    dataframes = {'periodos': dados_pivot_periodos,
                  'dias': dados_pivot_dias}
    

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

# --- Análises de Gráficos ---
      
def analisarFigura1(llm, document, client_name, dataframes):
            
    document.add_paragraph(f"A figura abaixo nos dá uma visão geral sobre quem são os melhores concorrentes da {client_name}, "
                        "segundo os indicadores seguidores, seguindo, quantidade de posts e de contagem de hashtags, "
                        "tanto a nível de perfil quanto a nível de publicações.")
        
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_1_Figura1(client_name, dataframes['df_profiles_posts'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))

    # Gerar Análise dos Dados da Figura 1 
    textos_analises = []
                
    prompt = f"""
        Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Produza uma análise exclusivamente descritiva detalhada dos dados do gráfico de barras abaixo. 
        Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
        Formato: Responda apenas o parágrafo da análise.
        Requisito: Inicie o texto dizendo "Segundo os gráficos acima...
        
        Dados dos Top 10 concorrentes mais seguidores: {dict_df['followers']}

        Dados dos Top 10  concorrentes mais seguindo pessoas: {dict_df['follows']}

        Dados dos Top 10  concorrentes com mais publicações: {dict_df['posts_count']}
    
    """
            
    analise = llm.invoke(prompt)
    document.add_paragraph(analise.content.replace('\n',''))  
            
    bio_cols = ['username', 'fullName', 'biography', 'externalUrl']        
    usernames = dict_df['followers']['username'].unique()
    profiles_df = dataframes['df_profiles_posts']
    df_bios = profiles_df[profiles_df['username'].isin(usernames)][bio_cols]
    
    # Gerar Recomendações            
    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em posicionamento de marca. 
        Contexto: Com base nos dados abaixo, faça uma análise do posicionamento dos melhores concorrentes da empresa cliente "{client_name}" no instagram.
        Tarefa: Gere um texto detalhado de 1 parágrafo com sua análise e que responda as perguntas: as biografias estão claras? Explicam o que as marcas fazem? Usam CTAs ou emojis estratégicos?. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Com base nas análises das biuografias dos concorrentes, foi possível perceber...". 
        
        Biografias: {df_bios}
    
        """
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')])

def analisarFigura2(llm, document, dataframes):
    document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre os concorrentes "
                        "por meio dos resultados de uma análise de clusterização. Esta análise é um tipo de análise estatística que "
                        "tem o objetivo de encontrar padrões ocultos em conjuntos de dados. Com ela, será possível perceber "
                        "quais concorrentes tem comportamentos parecidos no que diz respeito a Seguidores, Seguindo e Quantidade de Posts, sendo possível assim "
                        " a segmentação dos concorrentes em poucos grupos com alto grau de similaridade entre si.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_1_Figura2(dataframes['df_profiles_posts'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))

    # Gerar Análise dos Dados da Figura 1
    textos_analises = []
    for nome_df, df in dict_df.items():
            
        if nome_df == 'df_original_copy':
            inicio = "De acordo com o primeiro gráfico da figura acima..."
        elif nome_df == 'df_media_1':
            inicio = "Já de acordo com o segundo gráfico da figura..."
        elif nome_df == 'df_media_2':
            inicio = "Quanto ao último gráfico da figura acima..."
        elif nome_df == 'df_media_3':
            inicio = "Já no último gráfico da figura acima..."
                
        prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise descritiva detalhada com a carcaterização e personificação de cada um dos clusters gerados da clusterização de concorrentes do cliente "{client_name}", segundo os dados abaixo. 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df}
        """
            
        analise = llm.invoke(prompt)
        textos_analises.append(analise.content)
            
    analise_figura_1 = ' '.join(textos_analises)
    document.add_paragraph('\n'.join(textos_analises))

    prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia de conteúdo do cliente "{client_name}" no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise_figura_1}
    """
            
    print()
    recomendacoes = llm.invoke(prompt)
    print(f'Recomendações: {recomendacoes.content}')
    document.add_paragraph(recomendacoes.content) 
            
def analisarFigura3(llm, client_name, document, dataframes):
    document.add_paragraph(f" Na análise seguinte, será possível perceber uma visão geral sobre as hashtags utilizadas pelos concorrentes "
                                "por meio de uma nuvem de palavras. Uma nuvem de palavras (ou word cloud) é uma representação visual de texto onde "
                                "as palavras mais frequentes aparecem em destaque, com um tamanho maior ou cor diferente, enquanto as menos comuns "
                                "são menores. É usada para identificar rapidamente os termos mais importantes ou populares em um conjunto de dados textuais.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, df = Secao_2_1_Figura3(dataframes['posts_df'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))

    # Gerar Análise dos Dados da Figura 1 
    prompt = f"""
        Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Produza uma análise descritiva detalhada que responda a pergunta: Quais são as hashtags mais e menos utilizadas pelos concorrentes da "{client_name}" no instagram?. 
        Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
        Formato: Responda apenas os parágrafos da análise.
        Requisito: Inicie o texto dizendo "De acordo com a nuvem de palavras acima...".
        Dados: {df}
    """
            
    analise = llm.invoke(prompt)
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Com base nesta analise abaixo, quais hashtags você sugere que o cliente "{client_name}" utilize no instagram? Justifique de forma clara suas recomendações segundo os dados.
        Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Com base nas análises acima..."
        Analises: {analise.content.replace('\n','')}
    """

    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')])
  
def analisarFigura4(llm, client_name, document, dataframes):
            
    document.add_paragraph(f"A figura abaixo nos dá uma visão geral sobre quem são os melhores concorrentes do negócio, "
                                "segundo os indicadores de engajamento, tanto a nível de perfil quanto a nível de publicações.")
        
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_2_Figura4(client_name, dataframes['posts_df'], dataframes['df_profiles_posts'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
                
    prompt_perfis = f"""
            Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
            Contexto: Analise os dados abaixo e responda a pergunta: quem são os concorrentes com a maior quantidade de likes, comentários e taxa de engajamento do cliente "{client_name}"?. 
            Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua resposta. 
            Formato: Responda apenas o parágrafo da resposta.
            Requisito: Inicie o texto dizendo "De acordo com os 3 primeiros gráficos da figura acima...".
            
            Top 10 perfis com maior quantidade de likes: {dict_df['top_10_likes_profiles']}

            Top 10 perfis com maior quantidade de comments: {dict_df['top_10_comments_profiles']}

            Top 10 perfis com maior quantidade de percentual de engajamento: {dict_df["top_10_perc_engaj_profiles"]}
    """
            
    analise_perfis = llm.invoke(prompt_perfis)            
    document.add_paragraph(analise_perfis.content.replace('\n',''))

    prompt_posts = f"""
            Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
            Contexto: Analise os dados abaixo e responda a pergunta: quais são as publicações dos concorrentes com a maior quantidade de likes, comentários e taxa de engajamento do cliente "{client_name}"?. 
            Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua resposta. 
            Formato: Responda apenas o parágrafo da resposta.
            Requisito: Inicie o texto dizendo "Já de acordo com os 3 últimos gráficos da figura acima...".
            
            Top 10 perfis com maior quantidade de likes: {dict_df['top_10_likes_posts']}

            Top 10 perfis com maior quantidade de comments: {dict_df['top_10_comments_posts']}

            Top 10 perfis com maior quantidade de percentual de engajamento: {dict_df['top_10_engaj_posts']}
    """
            
    analise_posts = llm.invoke(prompt_posts)            
    document.add_paragraph(analise_posts.content.replace('\n',''))

    posts_df = dataframes['posts_df']

    list_posts = dict_df['top_10_engaj_posts']['shortCode'].unique()
    cols_content = ['shortCode', 'caption', 'hashtags', 'mentions', 'firstComment', 'latestComments']
    df_content = posts_df[posts_df['shortCode'].isin(list_posts)][cols_content]
    
    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em análise de conteúdo. 
        Contexto: Analise o melhores conteúdos dos concorrentes da empresa "{client_name}" abaixo e responda as perguntas: Qual extamanete é o conteúdo das publicações? O que as pessoas comentaram sobre o contéudo?  
        Tarefa: Gere um texto detalhado de 1 parágrafo com suas análises. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Quanto ao conteúdo das publicações,..."
        Publicações: {df_content}
    """
            
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise_perfis.content.replace('\n',''), analise_posts.content.replace('\n',''), recomendacoes.content.replace('\n','')])

def analisarFigura5(llm, client_name, document, dataframes):
    
    document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre quais os formatos de conteúdo "
                                "os concorrentes mais utilizam, bem como os que mais geraram engajamento.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_2_Figura5(dataframes['dados_pivot_count'], dataframes['dados_pivot_total'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
                
    prompt = f"""
        Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Analise os dados abaixo e responda as perguntas: Quais são os formatos mais utilizados e que mais geram engajamento para os concorrentes da "{client_name}"? 
        Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
        Formato: Responda apenas o parágrafo da análise.
        Requisito: Inicie o texto dizendo "No gráfico acima, é possível perceber que...".
        
        Quantidade de posts por tipo dos 3 melhores concorrentes: {dict_df['Type']}
        
        Quantidade de Interações por tipo dos 3 melhores concorrentes: {dict_df['total']}
    
    """
            
    analise = llm.invoke(prompt)
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Com base nesta analise abaixo, quais formatos de conteúdo a empresa "{client_name}" deve explorar no instagram? Justifique de forma clara cada uma delas
        Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Com base nas análises acima..."
        Analises: {analise.content.replace('\n','')}
    """
            
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content)

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')]) 

def analisarFigura6(llm, client_name, document, dataframes):
    
    document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre quais os formatos de conteúdo "
                                "mais geraram curtidas e comentários para os concorrentes.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_2_Figura6(dataframes['dados_pivot_likes'], dataframes['dados_pivot_comments'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
                
    prompt = f"""
        Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Produza uma análise detalhada que responda as perguntas: quais foram os tipos de conteúdo que mais geraram likes e comentários para os 3 melhores concorrentes? 
        Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
        Formato: Responda apenas o parágrafo da análise.
        Requisito: Inicie o texto dizendo "De acordo com os gráficos acima...".
        
        Proporção de comentários dos 3 melhores concorrentes por tipo de publicação: {dict_df['comments']}
        
        Proporção de likes dos 3 melhores concorrentes por tipo de publicação: {dict_df['likes']}
    """
            
    analise = llm.invoke(prompt)            
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Com base nesta analise abaixo, quais os tipos de conteúdo você sugere para a empresa "{client_name}" no instagram para gerar engajamento? Justifique de forma clara cada uma delas
        Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Com base nas análises acima..."
        Analises: {analise.content.replace('\n','')}
    """
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')]) 

def analisarFigura7(llm, client_name, document, dataframes):
            
    document.add_paragraph(f"A figura abaixo nos dá uma visão geral sobre quem são os melhores concorrentes do negócio, "
                                "segundo os indicadores de engajamento, tanto a nível de perfil quanto a nível de publicações.")
        
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_3_Figura7(client_name, dataframes['periodo_df'], dataframes['dias_df'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
    
    prompt = f"""
        Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Produza uma análise detalhada dos dados abaixo dos concorrentes da empresa "{client_name}" e responda a pergunta: Quais periodos e dias da semana os concorrentes mais publicam? 
        Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
        Formato: Responda apenas o parágrafo da análise.
        Requisito: Inicie o texto dizendo "De acordo com os gráficos acima...".
        
        Quantidade de publicações pelo periodo do dia: {dict_df['periodos']}
        
        Quantidade de publicações pelo dia da semana: {dict_df['dias']}
    """
            
    analise = llm.invoke(prompt)            
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Com base nesta analise abaixo, os melhores dias e periodos para a empresa "{client_name}" publicar no instagram? Justifique de forma clara cada uma delas
        Tarefa: Gere um texto detalhado de 1 parágrafo com sua resposta. 
        Formato: Responda apenas o parágrafo das respostas.
        Requisito: Inicie o texto dizendo "Com base nas análises acima..."
        Analises: {analise.content.replace('\n','')}
    """
            
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')])

def analisarFigura8(llm, client_name, document, dataframes):
    
    document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre a proporção de formatos pelo dia da semana, "
                                "com o objetivo de compreender em qual proporção os concorrentes publicam cada um dos tipos de conteúdos.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_3_Figura8(dataframes['dados_pivot_periodos'], dataframes['dados_pivot_dias'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
                
    prompt = f"""
            Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
            Contexto: Produza uma análise detalhada que responda as perguntas: quais os tipos de conteúdos mais postados pelos concorrentes nos dias da semana?. 
            Tarefa: Gere um texto detalhado de 1 parágrafo com sua resposta. 
            Formato: Responda apenas o parágrafo da análise.
            Requisito: Inicie o texto dizendo 'De acordo com o gráfico acima...'.
    
            Quantidade de posts por tipo e dias da semana: {dict_df['dias']}

    """
            
    analise = safe_invoke(llm, prompt)        
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, qual a proporção de tipos de conteúdo nos dias da semana que a empresa "{client_name}" deveria ter no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise.content.replace('\n','')}
            """
            
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')]) 

def analisarFigura9(llm, client_name, document, dataframes):
            
    document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre quais os formatos de conteúdo "
                                "mais publicados por periodo, com o objetivo se se compreender esta relação.")
            
    # Adiciona Figura 1
    paragrafo_da_imagem = document.add_paragraph()
    paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_da_imagem = paragrafo_da_imagem.add_run() 
    chart_buffer, dict_df = Secao_2_3_Figura9(dataframes['dados_pivot_periodos'], dataframes['dados_pivot_dias'])
    run_da_imagem.add_picture(chart_buffer, width=Inches(6))
                
    prompt = f"""
            Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
            Contexto: Produza uma análise detalhada e reponda as perguntas: quais os tipos de conteúdos mais postados pelos concorrentes nos periodos? 
            Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
            Formato: Responda apenas o parágrafo da análise.
            Requisito: Inicie o texto dizendo "Segundo o gráfico acima...".
            
            Dados: {dict_df['periodos']}
    """
            
    analise = safe_invoke(llm, prompt)
    document.add_paragraph(analise.content.replace('\n',''))

    prompt = f"""
        Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
        Contexto: Com base nesta analise abaixo, qual a proporção de tipos de conteúdo nos periodos que a empresa "{client_name}" deveria ter no instagram? Justifique de forma clara sua resposta
        Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
        Formato: Responda apenas o parágrafo das recomendações.
        Requisito: Inicie o texto dizendo "Com base nas análises acima..."
        Analises: {analise.content.replace('\n','')}
    """
            
    recomendacoes = llm.invoke(prompt)
    document.add_paragraph(recomendacoes.content.replace('\n',''))

    return '\n'.join([analise.content.replace('\n',''), recomendacoes.content.replace('\n','')]) 

# --- Funções de Formatação ---

def definir_estilos(document):
    """
    Centraliza todas as modificações de estilo do documento para garantir consistência.
    """
    # 1. ALTERAR A FONTE PADRÃO (ESTILO 'NORMAL')
    estilo_normal = document.styles['Normal']
    fonte_normal = estilo_normal.font
    fonte_normal.name = 'Abadi'
    fonte_normal.size = Pt(11)

    # 2. ALTERAR ESTILO DO TÍTULO PRINCIPAL (NÍVEL 1)
    # Usado pela função `gerarCapa` e outros títulos de nível 1
    estilo_h1 = document.styles['Heading 1']
    fonte_h1 = estilo_h1.font
    fonte_h1.name = 'Abadi'
    fonte_h1.size = Pt(22)
    fonte_h1.bold = True
    fonte_h1.color.rgb = RGBColor(0x0A, 0x25, 0x40)
    paragrafo_h1 = estilo_h1.paragraph_format
    paragrafo_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo_h1.space_before = Pt(12)
    paragrafo_h1.space_after = Pt(12)

    # 3. ALTERAR ESTILO DOS SUBTÍTULOS (NÍVEL 2)
    estilo_h2 = document.styles['Heading 2']
    fonte_h2 = estilo_h2.font
    fonte_h2.name = 'Abadi'
    fonte_h2.size = Pt(14)
    fonte_h2.bold = True
    fonte_h2.color.rgb = RGBColor(0x0A, 0x25, 0x40)  # Azul corporativo
    paragrafo_h2 = estilo_h2.paragraph_format
    paragrafo_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo_h2.space_before = Pt(18)
    paragrafo_h2.space_after = Pt(6)

    # 4. ALTERAR ESTILO DOS TÍTULOS DE TÓPICOS (NÍVEL 3)
    estilo_h3 = document.styles['Heading 3']
    fonte_h3 = estilo_h3.font
    fonte_h3.name = 'Abadi'
    fonte_h3.size = Pt(12)
    fonte_h3.bold = True
    fonte_h3.color.rgb = RGBColor(0x0A, 0x25, 0x40)
    paragrafo_h3 = estilo_h3.paragraph_format
    paragrafo_h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo_h3.space_before = Pt(10)
    paragrafo_h3.space_after = Pt(4)

def gerarCapa(document, titulo, cliente, autor, data):
    """
    Gera a capa do documento Word. A formatação do título vem do estilo 'Heading 1'.
    """
    # 1. Adicionar a logo da empresa na capa
    p_logo = document.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    # CORREÇÃO: Usando um valor sensato para a largura, como 6 cm.
    run_logo.add_picture(str(settings.LOGO_PATH), width=Cm(6))

    # 2. Título do Relatório (usa o estilo 'Heading 1' definido anteriormente)
    document.add_heading(titulo, level=1)

    # 3. Informações do Cliente e Autor
    document.add_paragraph() # Espaçamento
    p_info = document.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f'Preparado para:\n').bold = True
    p_info.add_run(f'{cliente}\n\n')
    p_info.add_run(f'Análise por:\n').bold = True
    p_info.add_run(f'{autor}\n\n')

    # 4. Data
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.add_run(data)
    document.add_page_break()

# --- Função Principal de Geração de Relatório ---
def generate_full_report(llm, dataframes, client_name, output_path, template_path):
    
    """Gera o relatório completo em .docx."""
   
    # Criar Documento
    document = Document(settings.TEMPLATE_PATH)

    definir_estilos(document)

    # Definir Informações do Cliente
    titulo_analise = "Análise de Concorrentes no Instagram"
    nome_cliente = client_name
    nome_autor = "Equipe AI Social"
    data_analise = f"{date.today().strftime("%A, %d de %B de %Y")}"

    # Gerar Relatório:
    
    # Capa e Resumo
    gerarCapa(document, titulo_analise, nome_cliente, nome_autor, data_analise) 
    
    # 1. Introdução
    document.add_heading("1.0 Introdução", level=2)
    intro = ('Este relatório apresenta uma análise competitiva aprofundada do Instagram, '
         'utilizando uma metodologia baseada em dados para desvendar as táticas e o '
         'desempenho dos concorrentes. A pesquisa visa transformar dados brutos em '
         'inteligência acionável, aprofundando-se no que impulsiona o engajamento, '
         'quais formatos e temas de conteúdo geram maior interação, e quais os '
         'melhores horários para publicação. ')
    document.add_paragraph(intro) 

    # 2. Analise de Concorrentes
    document.add_heading("2.0 Análise dos Concorrentes", level=2)
    texto_secao_2_1 = (f"Nesta seção será realizada uma análise comparativa entre os concorrentes do {nome_cliente}, "
            "a fim de traçar os seus perfis. Além de serem analisadas métricas de performance, frequência e recência, "
            "também serão analisados, qualitativamente, seus respectivos conteúdos, bem como o tom de voz, tópicos frequentes "
            "e posicionamento de marca. ")       
    
    document.add_paragraph(texto_secao_2_1)
    
    # 2.1 Análise de Perfil dos Concorrentes
    document.add_heading("2.1 Análise de Perfil dos Concorrentes", level=3)
    analises_figura_1 = analisarFigura1(llm, document, client_name, dataframes)
    analises_figura_3 = analisarFigura3(llm, client_name, document, dataframes)
    
    # 2.2 Análise de Engajamento por Postagem
    texto_secao_2_1 = (f"Nesta seção será realizada uma análise comparativa entre os as publicações dos concorrentes do {nome_cliente}, "
            "a fim de compreender suas respectivas estratégias de conteúdo. Além de serem analisadas métricas como curtidas e comentários "
            "também serão analisados, qualitativamente, seus respectivos conteúdos, bem como o tom de voz, tópicos frequentes "
            "e posicionamento de marca.")
    document.add_heading("2.2 Análise de Engajamento por Postagem", level=3)
    document.add_paragraph(texto_secao_2_1)  
    analises_figura_4 = analisarFigura4(llm, client_name, document, dataframes)
    analises_figura_5 = analisarFigura5(llm, client_name, document, dataframes)
    analises_figura_6 = analisarFigura6(llm, client_name, document, dataframes)
    time.sleep(60) 
    
    # 2.3 Frequência e Consistência de Publicação
    texto_secao_2_1 = (f"Nesta seção será realizada uma análise temporal das publicações dos concorrentes do {nome_cliente}, "
                    "a fim de compreender suas respectivas estratégias de conteúdo, mais especificamente, quais os melhores horários, "
                    "periodos e dias para publicar no feed.")
    document.add_heading("2.3 Frequência e Consistência de Publicação", level=3)
    document.add_paragraph(texto_secao_2_1)
    analises_figura_7 = analisarFigura7(llm, client_name, document, dataframes)
    analises_figura_8 = analisarFigura8(llm, client_name, document, dataframes)
    analises_figura_9 = analisarFigura9(llm, client_name, document, dataframes)

    
    analises = '\n'.join([analises_figura_1, analises_figura_3, analises_figura_4, 
                          analises_figura_5, analises_figura_6, analises_figura_7,
                          analises_figura_8, analises_figura_9])
    
    # 3. Recomendações Gerais
    document.add_heading("3.0 Conclusões\Recomendações Finais", level=2)
    
    prompt = f""" 
            Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
            Contexto: Com base em todas essas analises, gere um conclusão e resuma todas as recomendações dadas a empresa "{client_name}". 
            Tarefa: Gere um texto detalhado de 1 parágrafo com suas conclusões e recomendações finais. 
            Formato: Responda apenas o parágrafo da análise.
            Requisito: Inicie o texto dizendo "Com base em todas as análises realizadas...".
            
            Análises: {analises}
             
    """ 
    
    conclusao = llm.invoke(prompt)
    document.add_paragraph(conclusao.content.replace('\n',''))

    # Salva o documento
    document.save(output_path)  
