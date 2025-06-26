import io
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import gridspec
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from langchain_ollama import ChatOllama
from .AutoClusterHPO import AutoClusterHPO
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from wordcloud import WordCloud
from datetime import date
import locale
import numpy as np

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil')
    except locale.Error:
        # Caso nenhum dos locales seja encontrado, imprime um aviso.
        # A data pode sair em inglês ou no formato padrão do sistema.
        print("Aviso: Não foi possível configurar o idioma para português do Brasil. A data pode ser exibida em outro idioma.")

# --- Funções Auxiliares de Estilização ---
def set_cell_shading(cell, hex_color: str):
    """Define a cor de fundo de uma célula da tabela.""" 
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading_elm) 
    except Exception as e:
        print(f"Não foi possível aplicar a cor de fundo à célula: {e}") 

def set_header_style(table):
    """Aplica cor de fundo cinza e negrito ao cabeçalho da tabela.""" 
    header_cells = table.rows[0].cells
    for cell in header_cells:
        set_cell_shading(cell, 'D9D9D9') # Cinza claro
        
        # Para aplicar negrito
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run()
        run.text = cell.text
        run.bold = True
        p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER

# --- Funções de Geração de Gráficos ---
def Secao_2_1_Figura1(client_name, profile_df, posts_df):

    def tratarDados():
        
        # LINHA CRÍTICA A SER ADICIONADA:
        # Carregue o arquivo JSON para um DataFrame
        posts_df = pd.read_json("data/raw/post_data.json")
        profile_df = pd.read_json("data/raw/profile_data.json")

        # Tratar Dados
        posts_df['data_hora'] = pd.to_datetime(posts_df['timestamp'])

        posts_df_gruped = posts_df.groupby(['ownerId', 'ownerUsername']).agg(
        commentsSum=('commentsCount', 'sum'),
        likesSum=('likesCount', 'sum'),
        minData=('data_hora', 'min'),
        maxData=('data_hora', 'max'),
        count=('ownerId', 'count')
        ).reset_index()

        df_profiles_posts_int = pd.merge(profile_df, posts_df_gruped, left_on='id', right_on='ownerId', how='left').drop(['ownerId'], axis=1)
        df_profiles_posts_int['TOTAL ENGAJAMENTO'] = (df_profiles_posts_int['commentsSum'] + df_profiles_posts_int['likesSum'])
        df_profiles_posts_int[r'% ENGAJAMENTO'] =  df_profiles_posts_int['TOTAL ENGAJAMENTO'] / df_profiles_posts_int['followersCount']
        df_profiles_posts_int['RECENCIA'] = 1 / ((df_profiles_posts_int['maxData'].max() - df_profiles_posts_int['maxData']).dt.days + 1)
        df_profiles_posts_int['FREQUENCIA'] = df_profiles_posts_int['count'] / ((df_profiles_posts_int['maxData'] - df_profiles_posts_int['minData']).dt.days + 1)
        
        return df_profiles_posts_int
       
    def plotarBarraMax(client_name, x_col, y_col, ax1):
        
        # Calcular Estatísticas
        top_10 = df_profiles_posts_int.groupby([y_col])[x_col].max().sort_values(ascending=True).tail(10).reset_index()
        
        # --- 4. Configuração do Primeiro Gráfico (Superior) ---
        # Preparar cores e rótulos para o primeiro gráfico
        cores1 = ['#1f77b4'] * 10 # Cor padrão para as barras
        
        if client_name in list(top_10['username']):
            try:
                # Encontra a posição (índice) do elemento X nos dados ordenados
                indice_x = top_10.loc[top_10[y_col] == client_name].index[0]
                cores1[indice_x] = '#000080' # Cor de destaque para o elemento X
                print('Cor alterada com Sucesso!')
            except KeyError:
                indice_x = -1 # Trata o caso do elemento não ser encontrado

        # Plota o gráfico de barras
        barras1 = ax1.barh(top_10[y_col], top_10[x_col], color=cores1)
        ax1.bar_label(barras1, fmt='%d', padding=5)
        ax1.set_title(f'{y_col} X {x_col}')
        ax1.set_xlabel(x_col)
        ax1.set_ylabel(y_col)

        return top_10

    # --- 3. Criação da Figura e dos Gráficos (Subplots) ---
    # Cria a figura com 2 linhas e 1 coluna de gráficos
    fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(16, 5))

    df_profiles_posts_int = tratarDados()
      
    top_10_followers = plotarBarraMax(client_name, 'followersCount', 'username', ax1)
    top_10_follows = plotarBarraMax(client_name, 'followsCount', 'username', ax2)
    top_10_posts_count = plotarBarraMax(client_name, 'postsCount', 'username', ax3)

    dataframes = {
                    "followers": top_10_followers,
                    "follows": top_10_follows,
                    "posts_count": top_10_posts_count
    }

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # Para exibir a figura diretamente (se estiver em um ambiente interativo como Jupyter)
    plt.close()

    buffer.seek(0)

    print("Figura 'graficos_de_barras_destacados.png' gerada com sucesso.")

    return buffer, dataframes

def Secao_2_1_Figura2(profile_df, posts_df):
    
    def tratarDados():
        
        # LINHA CRÍTICA A SER ADICIONADA:
        posts_df = pd.read_json("data/raw/post_data.json")
        profile_df = pd.read_json("data/raw/profile_data.json")

        # Tratar Dados
        posts_df['data_hora'] = pd.to_datetime(posts_df['timestamp'])

        posts_df_gruped = posts_df.groupby(['ownerId', 'ownerUsername']).agg(
        commentsSum=('commentsCount', 'sum'),
        likesSum=('likesCount', 'sum'),
        minData=('data_hora', 'min'),
        maxData=('data_hora', 'max'),
        count=('ownerId', 'count')
        ).reset_index()

        df_profiles_posts_int = pd.merge(profile_df, posts_df_gruped, left_on='id', right_on='ownerId', how='left').drop(['ownerId'], axis=1)
        df_profiles_posts_int['TOTAL ENGAJAMENTO'] = (df_profiles_posts_int['commentsSum'] + df_profiles_posts_int['likesSum'])
        df_profiles_posts_int[r'% ENGAJAMENTO'] =  df_profiles_posts_int['TOTAL ENGAJAMENTO'] / df_profiles_posts_int['followersCount']
        df_profiles_posts_int['RECENCIA'] = 1 / ((df_profiles_posts_int['maxData'].max() - df_profiles_posts_int['maxData']).dt.days + 1)
        df_profiles_posts_int['FREQUENCIA'] = df_profiles_posts_int['count'] / ((df_profiles_posts_int['maxData'] - df_profiles_posts_int['minData']).dt.days + 1)
        
        return df_profiles_posts_int
    
    def plotarFiguraNColsPCA(df_original, df_cluster, fig):
            
        def plotarDispersao(fig):
            
            cluster_columns = [column for column in df_cluster.columns if column != 'Clusters (AutoClusterHPO)']
            
            # Redução de dimensionalidade com PCA
            pca = PCA(n_components=2)
            df_pca = pca.fit_transform(df_original[cluster_columns])
            df_original_copy['pca_x'] = df_pca[:, 0]
            df_original_copy['pca_y'] = df_pca[:, 1]
                
            # Visualização dos clusters com PCA
            axes_gs_principal = fig.add_subplot(gs_principal[0])
            sns.scatterplot(data=df_original_copy, x='pca_x', y='pca_y', hue='Clusters (AutoClusterHPO)', palette='deep', s=100, ax=axes_gs_principal) 
            axes_gs_principal.set_title(f" X ".join(cluster_columns))
            axes_gs_principal.set_xlabel('Componente Principal 1')
            axes_gs_principal.set_ylabel('Componente Principal 2')
            axes_gs_principal.grid(True)

            return df_original_copy[['pca_x', 'pca_y', 'Clusters (AutoClusterHPO)']]
        
        def plotarBarras(fig):
            
            df_original['Clusters (AutoClusterHPO)'] = 'Cluster ' + df_original['Clusters (AutoClusterHPO)'].astype(str)
                
            media_1 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[0])].mean().reset_index().sort_values(by=str(df_cluster.columns[0]), ascending=True)
            media_2 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[1])].mean().reset_index().sort_values(by=str(df_cluster.columns[1]), ascending=True)
            media_3 = df_original.groupby('Clusters (AutoClusterHPO)')[str(df_cluster.columns[2])].mean().reset_index().sort_values(by=str(df_cluster.columns[2]), ascending=True)
                
                
            # Plotar Gráfico de Barra 1
            axes_gs_0 = fig.add_subplot(gs[0])
            axes_gs_0.set_title(f'Clusters X {str(df_cluster.columns[0])}')
            barras1 = axes_gs_0.barh(media_1['Clusters (AutoClusterHPO)'], media_1[str(df_cluster.columns[0])])
            axes_gs_0.bar_label(barras1, fmt='%d', padding=-45)
            axes_gs_0.set_xticklabels([])
            axes_gs_0.set_xlabel(f'Média de {str(df_cluster.columns[0])}')

            # Plotar Gráfico de Barra 2
            axes_gs_1 = fig.add_subplot(gs[1])
            axes_gs_1.set_title(f'Clusters X {str(df_cluster.columns[1])}')
            barras2 = axes_gs_1.barh(media_2['Clusters (AutoClusterHPO)'], media_2[str(df_cluster.columns[1])])
            axes_gs_1.bar_label(barras2, fmt='%d', padding=5)
            axes_gs_1.set_xticklabels([])
            axes_gs_1.set_xlabel(f'Média de {str(df_cluster.columns[1])}')
                
            # Plotar Gráfico de Barra 3
            axes_gs1_0 = fig.add_subplot(gs1[0])
            axes_gs1_0.set_title(f'Clusters X {str(df_cluster.columns[2])}')
            barras3 = axes_gs1_0.barh(media_3['Clusters (AutoClusterHPO)'], media_3[str(df_cluster.columns[2])])
            axes_gs1_0.bar_label(barras3, fmt='%d', padding=-30)
            axes_gs1_0.set_xticklabels([])
            axes_gs1_0.set_xlabel(f'Média de {str(df_cluster.columns[2])}')

            return media_1, media_2, media_3
    
        # Copiar df original
        df_original_copy = df_original.copy()
            
        gs_principal = gridspec.GridSpec(1, 3, figure=fig, width_ratios=[2, 1, 1], wspace=0.2)
        gs = gs_principal[1].subgridspec(2, 1, hspace=0.4)
        gs1 = gs_principal[2].subgridspec(2, 1, hspace=0.4)
            
        df_original_copy = plotarDispersao(fig)
        media_1, media_2, media_3 = plotarBarras(fig)

        dataframes = {'df_original_copy' : df_original_copy, 
                    'df_media_1' : media_1,
                    'df_media_2' : media_2,
                    'df_media_3' : media_3}

        return dataframes

    df_profiles_posts_int = tratarDados()
    
    # Separar Bases 
    df_original = df_profiles_posts_int
    df_cluster = df_profiles_posts_int[['followersCount', 'followsCount', 'postsCount']].copy()

    # Inicializar e aplicar o AutoCluster
    autocluster_tool = AutoClusterHPO(max_evals_per_algo=100) 
    df_original['Clusters (AutoClusterHPO)'], model, config, score, algo_name = autocluster_tool.fit_predict(df_cluster)

    fig = plt.figure(figsize=(16, 5))
    
    dataframes = plotarFiguraNColsPCA(df_original, df_cluster, fig)

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
        
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)
        
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_1_Figura3(profile_df, posts_df):

    def plotarNuvemPalavras():

        posts_df = pd.read_json("data/raw/post_data.json")

        lista_unica = []

        for sublista in posts_df['hashtags']:
            for item in sublista:
                lista_unica.append(item)

        # Seu texto aqui
        texto = " ".join(lista_unica)

        # Criar o objeto WordCloud
        nuvem_palavras = WordCloud(width=800, height=400, background_color="white").generate(texto)
        
        return nuvem_palavras, pd.DataFrame(lista_unica).value_counts()

    nuvem_palavras, df = plotarNuvemPalavras()

    # Exibir a imagem gerada
    plt.figure(figsize=(16, 8))
    plt.imshow(nuvem_palavras, interpolation='bilinear')
    plt.axis("off") # Remove os eixos

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
        
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)
        
    plt.close()

    buffer.seek(0)

    return buffer, df 

def Secao_2_2_Figura4(client_name, profile_df, posts_df):

    def tratarDados():
        
        # LINHA CRÍTICA A SER ADICIONADA:
        # Carregue o arquivo JSON para um DataFrame
        posts_df = pd.read_json("data/raw/post_data.json")
        profile_df = pd.read_json("data/raw/profile_data.json")

        # Tratar Dados
        posts_df['data_hora'] = pd.to_datetime(posts_df['timestamp'])

        posts_df_gruped = posts_df.groupby(['ownerId', 'ownerUsername']).agg(
        commentsSum=('commentsCount', 'sum'),
        likesSum=('likesCount', 'sum'),
        minData=('data_hora', 'min'),
        maxData=('data_hora', 'max'),
        count=('ownerId', 'count')
        ).reset_index()

        df_profiles_posts_int = pd.merge(profile_df, posts_df_gruped, left_on='id', right_on='ownerId', how='left').drop(['ownerId'], axis=1)
        df_profiles_posts_int['TOTAL ENGAJAMENTO'] = (df_profiles_posts_int['commentsSum'] + df_profiles_posts_int['likesSum'])
        df_profiles_posts_int[r'% ENGAJAMENTO'] =  df_profiles_posts_int['TOTAL ENGAJAMENTO'] / df_profiles_posts_int['followersCount']
        df_profiles_posts_int['RECENCIA'] = 1 / ((df_profiles_posts_int['maxData'].max() - df_profiles_posts_int['maxData']).dt.days + 1)
        df_profiles_posts_int['FREQUENCIA'] = df_profiles_posts_int['count'] / ((df_profiles_posts_int['maxData'] - df_profiles_posts_int['minData']).dt.days + 1)
        
        return df_profiles_posts_int
       
    def plotarBarraSum(client_name, df, x_col, y_col, ax1, fmt):
        
        # Calcular Estatísticas
        top_10 = df.groupby([y_col])[x_col].sum().sort_values(ascending=True).tail(10).reset_index()
        
        # --- 4. Configuração do Primeiro Gráfico (Superior) ---
        # Preparar cores e rótulos para o primeiro gráfico
        cores1 = ['#1f77b4'] * 10 # Cor padrão para as barras
        
        if client_name in list(top_10[y_col]):
            try:
                # Encontra a posição (índice) do elemento X nos dados ordenados
                indice_x = top_10.loc[top_10[y_col] == client_name].index[0]
                cores1[indice_x] = '#000080' # Cor de destaque para o elemento X
                print('Cor alterada com Sucesso!')
            except KeyError:
                indice_x = -1 # Trata o caso do elemento não ser encontrado

        # Plota o gráfico de barras
        barras1 = ax1.barh(top_10[y_col], top_10[x_col], color=cores1)
        ax1.bar_label(barras1, fmt=fmt, padding=5)
        ax1.set_title(f'{y_col} X {x_col}')
        ax1.set_xlabel(x_col)
        ax1.set_ylabel(y_col)

        return top_10

    # --- 3. Criação da Figura e dos Gráficos (Subplots) ---
    # Cria a figura com 2 linhas e 1 coluna de gráficos
    fig, axes = plt.subplots(2, 3, figsize=(16, 5))

    df_profiles_posts_int = tratarDados()

    posts_df['TOTAL ENGAJAMENTO'] = posts_df['likesCount'] + posts_df['commentsCount'] 
      
    top_10_followers = plotarBarraSum(client_name, df_profiles_posts_int, 'likesSum', 'username', axes[0, 0], '%d')
    top_10_follows = plotarBarraSum(client_name, df_profiles_posts_int, 'commentsSum', 'username', axes[0, 1], '%d')
    top_10_posts_count = plotarBarraSum(client_name, df_profiles_posts_int, '% ENGAJAMENTO', 'username', axes[0, 2], '%.2f')
    top_10_followers = plotarBarraSum(client_name, posts_df, 'likesCount', 'shortCode', axes[1, 0], '%d')
    top_10_follows = plotarBarraSum(client_name, posts_df, 'commentsCount', 'shortCode', axes[1, 1], '%d')
    top_10_posts_count = plotarBarraSum(client_name, posts_df, 'TOTAL ENGAJAMENTO', 'shortCode', axes[1, 2], '%d')

    dataframes = {
                    "followers": top_10_followers,
                    "follows": top_10_follows,
                    "posts_count": top_10_posts_count
    }

    # --- 6. Finalização e Exibição/Salvamento da Figura ---
    plt.tight_layout(rect=[0, 0, 1, 0.96]) # Ajusta o layout para evitar sobreposição

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # Para exibir a figura diretamente (se estiver em um ambiente interativo como Jupyter)
    plt.close()

    buffer.seek(0)

    print("Figura 'graficos_de_barras_destacados.png' gerada com sucesso.")

    return buffer, dataframes

def Secao_2_2_Figura5(posts_df, profile_df):

    def tratarDados(posts_df, profile_df):

        def filtro(usernames):
            return (posts_df['ownerUsername'] == usernames[0]) | (posts_df['ownerUsername'] == usernames[1]) | (posts_df['ownerUsername'] == usernames[2])

        posts_df = pd.read_json("data/raw/post_data.json")
        profile_df = pd.read_json("data/raw/profile_data.json")
        
        filtro = filtro(['tarcisiogdf', 'romeuzemaoficial', 'eduardoleite'])

        posts_df_top_3 = posts_df[filtro]

        posts_df_top_3_merged = pd.merge(posts_df_top_3, profile_df, how='left', left_on='ownerUsername', right_on='username')

        posts_df_top_3_grouped = posts_df_top_3_merged.groupby(['ownerUsername', 'type']).agg(
            countType=('type', 'count'),
            followersMax=('followersCount', 'max'),
            likesSum=('likesCount', 'sum'),
            commentsSum=('commentsCount', 'sum')
        ).reset_index()

        posts_df_top_3_grouped['ENGAJAMENTO TOTAL'] = posts_df_top_3_grouped['commentsSum'] + posts_df_top_3_grouped['likesSum']

        dados_pivot_count = posts_df_top_3_grouped.pivot(index='ownerUsername', columns='type', values='countType')
        dados_pivot_total = posts_df_top_3_grouped.pivot(index='ownerUsername', columns='type', values='ENGAJAMENTO TOTAL')

        return dados_pivot_count, dados_pivot_total   

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(x_column, fontsize=12)
        ax.set_ylabel(y_column, fontsize=12)
        ax.set_title(f'{y_column} X {x_column}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    dados_pivot_count, dados_pivot_total = tratarDados(posts_df, profile_df)

    dataframes = {'Type': dados_pivot_count,
                  'total': dados_pivot_total}
    
    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(2, 1, figsize=(15, 10))
    
    plotarBarrasAgrupadas(dados_pivot_count, 'countType', 'ownerUsername', 'Type', axes[0])
    plotarBarrasAgrupadas(dados_pivot_total, 'ENGAJAMENTO TOTAL', 'ownerUsername', 'Type', axes[1])

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes

def Secao_2_2_Figura6(posts_df, profile_df):

    def tratarDados(posts_df, profile_df):

        def filtro(usernames):
            return (posts_df['ownerUsername'] == usernames[0]) | (posts_df['ownerUsername'] == usernames[1]) | (posts_df['ownerUsername'] == usernames[2])

        posts_df = pd.read_json("data/raw/post_data.json")
        profile_df = pd.read_json("data/raw/profile_data.json")
        
        filtro = filtro(['tarcisiogdf', 'romeuzemaoficial', 'eduardoleite'])

        posts_df_top_3 = posts_df[filtro]

        posts_df_top_3_merged = pd.merge(posts_df_top_3, profile_df, how='left', left_on='ownerUsername', right_on='username')

        posts_df_top_3_grouped = posts_df_top_3_merged.groupby(['ownerUsername', 'type']).agg(
            countType=('type', 'count'),
            followersMax=('followersCount', 'max'),
            likesSum=('likesCount', 'sum'),
            commentsSum=('commentsCount', 'sum')
        ).reset_index()

        dados_pivot_likes = posts_df_top_3_grouped.pivot(index='ownerUsername', columns='type', values='likesSum')
        dados_pivot_comments = posts_df_top_3_grouped.pivot(index='ownerUsername', columns='type', values='commentsSum')

        return dados_pivot_likes, dados_pivot_comments 

    def plotarBarrasAgrupadas(dados_pivot, x_column, y_column, group_column, ax):

        categorias_principais = dados_pivot.index
        subcategorias = dados_pivot.columns

        n_categorias_principais = len(categorias_principais)
        n_subcategorias = len(subcategorias)

        x = np.arange(n_categorias_principais)
        largura_barra = 0.25

        # 5. Loop para criar as barras para cada subcategoria (cada produto)
        for i, produto in enumerate(subcategorias):
            # Cálculo da posição de cada barra dentro do grupo.
            # O cálculo desloca cada conjunto de barras em relação ao centro do grupo (x).
            # O objetivo é centralizar o cluster de barras em torno do tick do eixo X.
            posicao = x - (largura_barra * n_subcategorias / 2) + (i * largura_barra) + (largura_barra / 2)
            
            # Pega os valores para o produto atual
            valores = dados_pivot[produto]
            
            # Cria as barras
            barra = ax.bar(posicao, valores, largura_barra, label=produto)

            # Adiciona os rótulos de valor no topo de cada barra para clareza
            ax.bar_label(barra, padding=3, fmt='%d') # fmt='%d' para mostrar como inteiro


        # 6. Adicionar rótulos, título, legenda e outros detalhes de formatação
        ax.set_xlabel(x_column, fontsize=12)
        ax.set_ylabel(y_column, fontsize=12)
        ax.set_title(f'{y_column} X {x_column}', fontsize=16)

        # Posiciona os rótulos da categoria principal (Anos) no centro dos grupos de barras
        ax.set_xticks(x)
        ax.set_xticklabels(categorias_principais)

        # Adiciona a legenda para identificar as cores das barras
        ax.legend(title=group_column, bbox_to_anchor=(1.02, 1), loc='upper left')

        # Adiciona uma grade horizontal para facilitar a leitura dos valores
        ax.yaxis.grid(True, linestyle='--', alpha=0.7)
        ax.set_axisbelow(True) # Coloca a grade atrás das barras

        # Remove as bordas superior e direita para um visual mais limpo
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    dados_pivot_likes, dados_pivot_comments = tratarDados(posts_df, profile_df)

    dataframes = {'likes': dados_pivot_likes, 
                  'comments': dados_pivot_comments}
    
    # 4. Criar a figura e os eixos do gráfico
    fig, axes = plt.subplots(2, 1, figsize=(15, 10))
    
    plotarBarrasAgrupadas(dados_pivot_likes, 'likes', 'ownerUsername', 'Type', axes[0])
    plotarBarrasAgrupadas(dados_pivot_comments, 'comments', 'ownerUsername', 'Type', axes[1])

    # Ajusta o layout para garantir que nada (como a legenda) seja cortado
    fig.tight_layout()

    buffer = io.BytesIO() 
    
    # Para salvar a figura em um arquivo
    plt.savefig(buffer, format='png', dpi=300)

    # 7. Exibir o gráfico
    plt.close()

    buffer.seek(0)

    return buffer, dataframes


# --- Funções de Geração de Seções
def gerarCapaResumo(document, titulo, cliente, autor, data, resumo_profissional):
    """
    Gera um documento Word com capa e resumo profissional para uma análise de concorrentes.

    Args:
        titulo (str): O título do relatório.
        cliente (str): O nome do cliente para quem a análise foi preparada.
        autor (str): O nome do autor ou da agência/consultoria.
        data (str): A data de publicação do relatório.
        resumo_profissional (str): O texto do resumo profissional.
    """
    # --- Seção da Capa ---

    # Título do Relatório
    titulo_principal = document.add_heading(titulo, level=1)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_principal.runs:
        run.font.size = Pt(24)
        run.bold = True

    # Adiciona espaçamento após o título
    document.add_paragraph()
    document.add_paragraph()

    # Informações do Cliente e Autor
    p_info = document.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f'Preparado para:\n{cliente}\n\n').bold = True
    p_info.add_run(f'Análise por:\n{autor}\n\n')

    # Data
    p_data = document.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.add_run(data)

    # Adiciona quebra de página
    document.add_page_break()

    # --- Seção do Resumo Profissional ---

    # Título do Resumo
    titulo_resumo = document.add_heading('Resumo Profissional', level=2)
    for run in titulo_resumo.runs:
        run.font.size = Pt(16)
        run.bold = True

    # Corpo do Resumo
    p_resumo = document.add_paragraph(resumo_profissional)
    p_resumo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Salva o documento
    nome_arquivo = f"Analise_Concorrentes_Instagram_{cliente.replace(' ', '_')}.docx"
    document.save(nome_arquivo)
    print(f"Relatório '{nome_arquivo}' gerado com sucesso!")

# --- Função Principal de Geração de Relatório ---
def generate_full_report(client_name, profile_df, posts_df, content_analysis, output_path, template_path):
    
    """Gera o relatório completo em .docx."""
        
    def criarSecao2_1(client_name, profile_df, posts_df, content_analysis, output_path, template_path):
        
        def analisarFigura1():
            
            document.add_paragraph(f"A figura abaixo nos dá uma visão geral sobre quem são os melhores concorrentes do negócio, "
                                "segundo os indicadores de interesse, tanto a nível de perfil quanto a nível de publicações.")
        
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, dict_df = Secao_2_1_Figura1(client_name, profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.1:8b", temperature=0) 
            textos_analises = []
            for nome_df, df in dict_df.items():

                if not df.empty:
                    df_ord = df.sort_values(by=str(df.columns[-1]), ascending=False)
                else:
                    df_ord = df
                print(df_ord, end='\n')
            
                if nome_df == 'followers':
                    inicio = "De acordo com o primeiro gráfico da figura acima..."
                elif nome_df == 'follows':
                    inicio = "Já de acordo com o segundo gráfico da figura..."
                elif nome_df == 'posts_count':
                    inicio = "Quanto ao último gráfico da figura acima..."
                
                prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise exclusivamente descritiva detalhada dos dados do gráfico de barras abaixo. 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua descrição. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df_ord}
                """
            
                analise = llm.invoke(prompt)
                textos_analises.append(analise.content)
                document.add_paragraph(analise.content) 
            
            analise_figura_1 = ' '.join(textos_analises) 
            
            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia de conteúdo do cliente "{client_name}" no instagram? Justifique de forma clara suas recomendações segundo os dados.
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima, é recomendado...". Além disso, garanta que sua análise responda as perguntas estratégicas do negócio.
                Perguntas Estratégicas: Qual deve ser a meta mínima de seguidores se o cliente quiser estar em décimo entre os concorrentes mais seguidores? E se ele não quiser estar entre os 10 com mais posts e que mais seguem pessoas qual o numéro máximo desses indicadores?
                Analises: {analise_figura_1}
            """
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content)
            
        def analisarFigura2():
            document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre os concorrentes "
                                "por meio dos resultados de uma análise de clusterização. Esta análise é um tipo de análise estatística que "
                                "tem o objetivo de encontrar padrões ocultos em conjuntos de dados. Com ela, será possível perceber "
                                "quais concorrentes tem comportamentos parecidos no que diz respeito a Seguidores, Seguindo e Quantidade de Posts, sendo possível assim "
                                " a segmentação dos concorrentes em poucos grupos com alto grau de similaridade entre si.")
            
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, dict_df = Secao_2_1_Figura2(profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.2:latest", temperature=0) 
            textos_analises = []
            for nome_df, df in dict_df.items():
            
                if nome_df == 'df_original_copy':
                    inicio = "De acordo com o primeiro gráfico da figura acima..."
                elif nome_df == 'df_media_1':
                    inicio = "Já de acordo com o segundo gráfico da figura..."
                elif nome_df == 'df_media_2':
                    inicio = "Quanto ao último gráfico da figura acima..."
                elif nome_df == 'df_media_3':
                    inicio = "Já no último gráfico da figura acima..."
                
                prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise descritiva detalhada com a carcaterização e personificação de cada um dos clusters gerados da clusterização de concorrentes do cliente "{client_name}", segundo os dados abaixo. 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df}
                """
            
                analise = llm.invoke(prompt)
                textos_analises.append(analise.content)
            
            analise_figura_1 = ' '.join(textos_analises)
            document.add_paragraph('\n'.join(textos_analises))

            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia de conteúdo do cliente "{client_name}" no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise_figura_1}
            """
            
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content) 
            
        def analisarFigura3():
            document.add_paragraph(f" Na análise seguinte, será possível perceber uma visão geral sobre as hashtags utilizadas pelos concorrentes "
                                "por meio de uma nuvem de palavras. Uma nuvem de palavras (ou word cloud) é uma representação visual de texto onde "
                                "as palavras mais frequentes aparecem em destaque, com um tamanho maior ou cor diferente, enquanto as menos comuns "
                                "são menores. É usada para identificar rapidamente os termos mais importantes ou populares em um conjunto de dados textuais.")
            
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, df = Secao_2_1_Figura3(profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.2:latest", temperature=0) 

            prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise descritiva detalhada com a sua interpretação dos dados da tabela de frequências das hashtags utilizadas pelos concorrentes do cliente "{client_name}". 
                Tarefa: Gere um texto científico detalhado de 2 parágrafos com sua análise. 
                Formato: Responda apenas os parágrafos da análise.
                Requisito: Inicie o texto dizendo "De acordo com a nuvem de palavras acima...".
                Dados: {df}
                """
            
            analise = llm.invoke(prompt)
            document.add_paragraph(analise.content)

            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia de conteúdo do cliente "{client_name}" no instagram? Justifique de forma clara suas recomendações segundo os dados.
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise}
            """
            
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content)
    
        texto_secao_2_1 = (f"Nesta seção será realizada uma análise comparativa entre os concorrentes do Jeronim Rodrigues Ba, "
                    "a fim de traçar os seus perfis. Além de serem analisadas métricas de performance, frequência e recência, "
                    "também serão analisados, qualitativamente, seus respectivos conteúdos, bem como o tom de voz, tópicos frequentes "
                    "e posicionamento de marca. O principal objetivo desta seção é o de avaliar a presença visual, textual e social "
                    "dos concorrentes e seus respectivos posicionamentos com indicadores e análises estratégicos dos perfis e conteúdos "
                    "dos concorrentes.")
        
        document.add_heading("2.1 Análise de Perfil dos Concorrentes", level=2)
        document.add_paragraph(texto_secao_2_1)

        analisarFigura1()
        analisarFigura2()
        analisarFigura3()  

    def criarSecao2_2(client_name, profile_df, posts_df, content_analysis, output_path, template_path):
        
        def analisarFigura1():
            
            document.add_paragraph(f"A figura abaixo nos dá uma visão geral sobre quem são os melhores concorrentes do negócio, "
                                "segundo os indicadores de engajamento, tanto a nível de perfil quanto a nível de publicações.")
        
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, dict_df = Secao_2_2_Figura4(client_name, profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.2:latest", temperature=0) 
            textos_analises = []
            for nome_df, df in dict_df.items():

                if not df.empty:
                    df_ord = df.sort_values(by=str(df.columns[-1]), ascending=False)
                else:
                    df_ord = df
                print(df_ord, end='\n')
            
                if nome_df == 'followers':
                    inicio = "De acordo com o primeiro gráfico da figura acima..."
                elif nome_df == 'follows':
                    inicio = "Já de acordo com o segundo gráfico da figura..."
                elif nome_df == 'posts_count':
                    inicio = "Quanto ao último gráfico da figura acima..."
                
                prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise detalhada com a sua interpretação dos dados do gráfico de barras dos perfis concorrentes do cliente "{client_name}". 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df_ord}
                """
            
                analise = llm.invoke(prompt)
                textos_analises.append(analise.content)
            
            
            analise_figura_1 = ' '.join(textos_analises)
            document.add_paragraph('\n\n'.join(textos_analises))

            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia do cliente "{client_name}" no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise_figura_1}
            """
            
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content)
            
        def analisarFigura2():
            document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre quais os formatos de conteúdo "
                                "os concorrentes mais utilizam, bem como os que mais geraram engajamento.")
            
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, dict_df = Secao_2_2_Figura5(profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.2:latest", temperature=0) 
            textos_analises = []
            for nome_df, df in dict_df.items():
            
                if nome_df == 'Type':
                    inicio = "De acordo com o primeiro gráfico da figura acima..."
                elif nome_df == 'total':
                    inicio = "Já de acordo com o segundo gráfico da figura..."
                
                prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise detalhada com a sua interpretação dos dados dos gráficos dos perfis concorrentes do cliente "{client_name}" abaixo. 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df}
                """
            
                analise = llm.invoke(prompt)
                textos_analises.append(analise.content)
            
            analise_figura_1 = ' '.join(textos_analises)
            document.add_paragraph('\n\n'.join(textos_analises))

            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia do cliente "{client_name}" no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise_figura_1}
            """
            
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content) 

        def analisarFigura3():
            document.add_paragraph(f"Na análise que se segue, será possível perceber uma visão geral sobre quais os formatos de conteúdo "
                                "mais geraram curtidas e comentários para os concorrentes.")
            
            # Adiciona Figura 1
            paragrafo_da_imagem = document.add_paragraph()
            paragrafo_da_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_da_imagem = paragrafo_da_imagem.add_run() 
            chart_buffer, dict_df = Secao_2_2_Figura6(profile_df, posts_df)
            run_da_imagem.add_picture(chart_buffer, width=Inches(6))

            # Gerar Análise dos Dados da Figura 1
            llm = ChatOllama(model="llama3.2:latest", temperature=0) 
            textos_analises = []
            for nome_df, df in dict_df.items():
            
                if nome_df == 'likes':
                    inicio = "De acordo com o primeiro gráfico da figura acima..."
                elif nome_df == 'comments':
                    inicio = "Já de acordo com o segundo gráfico da figura..."
                
                prompt = f"""
                Persona: Você é um analista\estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Produza uma análise detalhada com a sua interpretação dos dados dos gráficos dos perfis concorrentes do cliente "{client_name}" abaixo. 
                Tarefa: Gere um texto científico detalhado de 1 parágrafo com sua análise. 
                Formato: Responda apenas o parágrafo da análise.
                Requisito: Inicie o texto dizendo {inicio}.
                Dados: {df}
                """
            
                analise = llm.invoke(prompt)
                textos_analises.append(analise.content)
            
            analise_figura_1 = ' '.join(textos_analises)
            document.add_paragraph('\n\n'.join(textos_analises))

            prompt = f"""
                Persona: Você é um estrategista de marketing de mídias sociais sênior, especialista em social metrics. 
                Contexto: Com base nesta analise abaixo, quais recomendações você sugere para a estratégia do cliente "{client_name}" no instagram? Justifique de forma clara cada uma delas
                Tarefa: Gere um texto detalhado de 1 parágrafo com suas recomendações. 
                Formato: Responda apenas o parágrafo das recomendações.
                Requisito: Inicie o texto dizendo "Com base nas análises acima..."
                Analises: {analise_figura_1}
            """
            
            print()
            recomendacoes = llm.invoke(prompt)
            print(f'Recomendações: {recomendacoes.content}')
            document.add_paragraph(recomendacoes.content) 


        texto_secao_2_1 = (f"Nesta seção será realizada uma análise comparativa entre os as publicações dos concorrentes do Jeronim Rodrigues Ba, "
                    "a fim de compreender suas respectivas estratégias de conteúdo. Além de serem analisadas métricas como curtidas e comentários "
                    "também serão analisados, qualitativamente, seus respectivos conteúdos, bem como o tom de voz, tópicos frequentes "
                    "e posicionamento de marca. O principal objetivo desta seção é o de avaliar os pontos fortes das melhores publicações "
                    "dos concorrentes e seus respectivos posicionamentos com indicadores e análises estratégicos dos perfis e conteúdos "
                    "dos concorrentes.")
        
        document.add_heading("2.2 Análise de Engajamento por Postagem", level=2)
        document.add_paragraph(texto_secao_2_1)  

        analisarFigura1()
        analisarFigura2()
        analisarFigura3()

    document = Document(template_path)

    titulo_analise = "Análise de Concorrentes no Instagram"
    nome_cliente = "Cliente X"
    nome_autor = "Sua Agência / Seu Nome"
    data_analise = f"{date.today().strftime("%A, %d de %B de %Y")}"
    texto_resumo = (
        "Este relatório apresenta uma análise competitiva do desempenho no Instagram, "
        "preparada para o Cliente X. O estudo foi conduzido entre maio e junho de 2025 e "
        "avaliou os perfis dos principais concorrentes: Concorrente A, Concorrente B e "
        "Concorrente C. A análise focou em métricas de engajamento, crescimento de seguidores, "
        "estratégia de conteúdo e frequência de postagens. Os resultados indicam que o "
        "Concorrente B lidera em engajamento, utilizando fortemente o formato de Reels com "
        "conteúdo gerado pelo usuário. O Concorrente A, por sua vez, demonstra um crescimento "
        "acelerado de seguidores através de parcerias com influenciadores. Identificou-se uma "
        "oportunidade para o Cliente X explorar o formato de Stories interativos (enquetes, "
        "caixas de perguntas), pouco utilizado pela concorrência. Recomenda-se a adoção de uma "
        "estratégia de conteúdo em vídeo mais robusta, similar à do Concorrente B, e a "
        "implementação de um calendário editorial que inclua colaborações estratégicas para "
        "aumentar o alcance e o engajamento do perfil."
    )
    
    gerarCapaResumo(document, titulo_analise, nome_cliente, nome_autor, data_analise, texto_resumo) 
    
    # Estrutura Inicial
    document.add_page_break()
    document.add_heading("1.0 Introdução", level=1)

    # Analise de Concorrentes
    document.add_heading("2.0 Análise dos Concorrentes", level=1)
    criarSecao2_1(client_name, profile_df, posts_df, content_analysis, output_path, template_path)
    criarSecao2_2(client_name, profile_df, posts_df, content_analysis, output_path, template_path)
    
    # Salva o documento
    document.save(output_path)  
